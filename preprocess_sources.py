# preprocess_sources.py
import os
import glob
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
import traceback
import chardet
from tqdm import tqdm  # Para barra de progresso
import sys
import unicodedata # Para normalização de texto (opcional aqui, mas pode ajudar)
import re # Para limpeza básica (opcional aqui)

# --- Configurações ---
SOURCE_DIR = r"C:\Users\pbm_s\OneDrive\Prova Digital"  # Pasta original
OUTPUT_DIR = os.path.join(SOURCE_DIR, "Processed_For_Index") # Novo diretório de saída
OCR_LANG = 'por'  # Idioma(s) para Tesseract (ex: 'por+eng' para português e inglês)
MIN_CHARS_PER_PAGE_THRESHOLD = 50  # Mínimo de caracteres digitais por página para NÃO usar OCR
OCR_DPI = 300  # Resolução para renderizar imagem para OCR
OVERWRITE_EXISTING = True # True para reprocessar e sobrescrever arquivos existentes no OUTPUT_DIR

# --- Configuração Opcional do Tesseract ---
# Descomente e ajuste se necessário (ex: Windows)
TESSERACT_CMD_PATH = r'C:\Program Files\Tesseract-OCR\tesseract-OCR.exe'
# try:
#     pytesseract.tesseract_cmd = TESSERACT_CMD_PATH
# except NameError:
#      pass # Ignora se TESSERACT_CMD_PATH não foi definido
# -----------------------------------------

# Garante que o diretório de saída exista
os.makedirs(OUTPUT_DIR, exist_ok=True)

def clean_text(text):
    """Limpeza básica e normalização unicode para o texto extraído."""
    if not text:
        return ""
    try:
        # Normalização Unicode
        text = unicodedata.normalize('NFKC', text)
        # Remover múltiplos espaços
        text = re.sub(r'\s+', ' ', text).strip()
        # Opcional: Outras limpezas podem ser adicionadas aqui se necessário
    except Exception as e:
        print(f"  [Aviso] Erro na limpeza básica de texto: {e}")
    return text

def process_pdf(pdf_path, output_txt_path):
    """
    Processa um PDF. Tenta extração digital, se falhar/insuficiente, tenta OCR.
    Salva o resultado como TXT. Retorna True se usou OCR, False caso contrário.
    """
    used_ocr = False
    try:
        doc = fitz.open(pdf_path)
        if doc.is_encrypted and not doc.authenticate(""):
            print(f"  [Aviso] PDF criptografado: {os.path.basename(pdf_path)}")
            with open(output_txt_path, "w", encoding="utf-8") as f_out:
                f_out.write(clean_text(f"[Erro: PDF criptografado: {os.path.basename(pdf_path)}]"))
            return False # Não usou OCR

        digital_texts = []
        total_digital_chars = 0
        needs_ocr_check = True

        # 1. Tenta extração digital e avalia necessidade de OCR
        for page_num, page in enumerate(doc):
            try:
                page_text = page.get_text("text", sort=True).strip()
                digital_texts.append(page_text)
                total_digital_chars += len(page_text)
            except Exception as e_extract:
                print(f"  [Erro] Extração digital página {page_num+1} de {os.path.basename(pdf_path)}: {e_extract}")
                digital_texts.append(f"[Erro extração pág {page_num+1}]")

        # Decide se OCR é necessário para o *documento inteiro*
        if doc.page_count > 0:
            avg_chars_per_page = total_digital_chars / doc.page_count
            if avg_chars_per_page >= MIN_CHARS_PER_PAGE_THRESHOLD:
                needs_ocr_check = False # Texto digital parece suficiente
                # print(f"  [Info] Texto digital suficiente para {os.path.basename(pdf_path)} (média: {avg_chars_per_page:.1f} chars/pág).")
        else:
             needs_ocr_check = False # Documento vazio

        # 2. Extrai o texto final (Digital ou OCR)
        final_texts = []
        if not needs_ocr_check:
            final_texts = digital_texts # Usa o texto digital já extraído
        else:
            print(f"  [Info] Texto digital insuficiente/ausente em {os.path.basename(pdf_path)}. Tentando OCR ({OCR_LANG})...")
            used_ocr = True # Marca que tentará OCR
            for page_num, page in enumerate(doc):
                page_text_ocr = f"[Erro OCR pág {page_num+1}]" # Default em caso de erro
                try:
                    # Renderiza página para imagem
                    pix = page.get_pixmap(dpi=OCR_DPI)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    # Executa OCR
                    page_text_ocr = pytesseract.image_to_string(img, lang=OCR_LANG, config='--psm 6').strip()
                except pytesseract.TesseractNotFoundError:
                     print(f"  [ERRO FATAL] Tesseract não encontrado! Verifique instalação/PATH/configuração.")
                     # Poderia sair do script aqui ou apenas parar de tentar OCR
                     sys.exit(1) # Sai do script se Tesseract não funciona
                except ImportError:
                     print(f"  [ERRO FATAL] Pillow ou Pytesseract não instalado!")
                     sys.exit(1)
                except Exception as e_ocr:
                    print(f"  [Erro] OCR página {page_num+1} de {os.path.basename(pdf_path)}: {e_ocr}")
                final_texts.append(page_text_ocr)

        # 3. Salva o resultado
        final_full_text = "\n\n".join(final_texts) # Junta páginas
        cleaned_final_text = clean_text(final_full_text) # Aplica limpeza básica

        with open(output_txt_path, "w", encoding="utf-8") as f_out:
            f_out.write(cleaned_final_text)

        doc.close()
        return used_ocr

    except Exception as e:
        print(f"  [ERRO GERAL] Processando PDF {os.path.basename(pdf_path)}: {e}")
        traceback.print_exc()
        # Tenta escrever um arquivo de erro
        try:
            with open(output_txt_path, "w", encoding="utf-8") as f_out:
                f_out.write(clean_text(f"[Erro geral processando PDF: {os.path.basename(pdf_path)} - {e}]"))
        except:
             pass # Ignora se nem o arquivo de erro puder ser escrito
        if 'doc' in locals() and doc:
            try: doc.close()
            except: pass
        return False # Assume que não usou OCR se deu erro geral


def process_docx(docx_path, output_txt_path):
    """Extrai texto de DOCX e salva como TXT."""
    try:
        doc = Document(docx_path)
        full_text_list = [p.text for p in doc.paragraphs]
        full_text = "\n".join(full_text_list)
        # Extração básica de tabelas
        try:
             for table in doc.tables:
                  full_text += "\n\n[INÍCIO TABELA]\n"
                  for row in table.rows:
                       full_text += "\t|\t".join(cell.text.strip() for cell in row.cells) + "\n"
                  full_text += "[FIM TABELA]\n\n"
        except Exception as e_table:
             # print(f"  [Aviso] Não extraiu tabela de {os.path.basename(docx_path)}: {e_table}")
             pass

        cleaned_text = clean_text(full_text)
        with open(output_txt_path, "w", encoding="utf-8") as f_out:
            f_out.write(cleaned_text)
    except Exception as e:
        print(f"  [ERRO] Processando DOCX {os.path.basename(docx_path)}: {e}")
        try:
            with open(output_txt_path, "w", encoding="utf-8") as f_out:
                f_out.write(clean_text(f"[Erro processando DOCX: {os.path.basename(docx_path)} - {e}]"))
        except: pass

def process_txt(txt_path, output_txt_path):
    """Lê TXT, detecta encoding (fallback UTF-8) e salva como UTF-8."""
    raw_text = ""
    try:
        # Detecta encoding
        detected_encoding = 'utf-8'
        try:
            with open(txt_path, "rb") as f_raw:
                 detected = chardet.detect(f_raw.read(30000)) # Aumenta bytes lidos
                 if detected and detected['encoding'] and detected['confidence'] > 0.6: # Limiar um pouco mais baixo
                     detected_encoding = detected['encoding']
                     # print(f"   [Debug] Detectado {detected_encoding} para {os.path.basename(txt_path)}")
        except Exception as e_chardet:
             # print(f"  [Aviso] Chardet falhou para {os.path.basename(txt_path)}: {e_chardet}. Usando UTF-8.")
             pass

        # Lê com encoding detectado ou fallback UTF-8
        try:
            with open(txt_path, "r", encoding=detected_encoding, errors='ignore') as f:
                raw_text = f.read()
        except LookupError: # Encoding inválido
            # print(f"  [Aviso] Encoding '{detected_encoding}' inválido. Tentando UTF-8.")
            with open(txt_path, "r", encoding='utf-8', errors='ignore') as f:
                raw_text = f.read()

        cleaned_text = clean_text(raw_text)
        with open(output_txt_path, "w", encoding="utf-8") as f_out:
            f_out.write(cleaned_text)

    except Exception as e:
        print(f"  [ERRO] Processando TXT {os.path.basename(txt_path)}: {e}")
        try:
            with open(output_txt_path, "w", encoding="utf-8") as f_out:
                f_out.write(clean_text(f"[Erro processando TXT: {os.path.basename(txt_path)} - {e}]"))
        except: pass

# --- Loop Principal ---
if __name__ == "__main__":
    print(f"Iniciando pré-processamento.")
    print(f"Pasta Fonte: {SOURCE_DIR}")
    print(f"Pasta Saída: {OUTPUT_DIR}")
    print(f"Modo OCR: {'Habilitado (se necessário)' if pytesseract else 'Desabilitado (pytesseract não importado)'}")
    print(f"Sobrescrever existentes: {OVERWRITE_EXISTING}")
    print("-" * 30)

    # Encontra todos os arquivos suportados na pasta fonte
    files_to_process = []
    supported_extensions = ["*.pdf", "*.docx", "*.txt"]
    for ext in supported_extensions:
        files_to_process.extend(glob.glob(os.path.join(SOURCE_DIR, "**", ext), recursive=True))

    print(f"Encontrados {len(files_to_process)} arquivos ({', '.join(supported_extensions)}) para verificar.")

    processed_count = 0
    ocr_count = 0
    error_count = 0

    # Itera pelos arquivos com barra de progresso
    for source_path in tqdm(files_to_process, desc="Processando arquivos"):
        try:
            # Define o caminho do arquivo de saída TXT
            relative_path = os.path.relpath(source_path, SOURCE_DIR)
            output_path_no_ext = os.path.join(OUTPUT_DIR, relative_path)
            # Adiciona extensão .txt ao nome original para evitar conflitos e rastrear origem
            output_txt_path = output_path_no_ext + ".txt"

            # Cria subdiretórios necessários na pasta de saída
            output_sub_dir = os.path.dirname(output_txt_path)
            os.makedirs(output_sub_dir, exist_ok=True)

            # Verifica se deve pular (se não for para sobrescrever e já existir)
            if not OVERWRITE_EXISTING and os.path.exists(output_txt_path):
                 # print(f"  [Skipping] Arquivo já existe: {os.path.basename(output_txt_path)}")
                 continue

            # Determina o tipo e processa
            ext = os.path.splitext(source_path)[1].lower()
            ocr_used_flag = False # Flag para contar OCR

            if ext == ".pdf":
                ocr_used_flag = process_pdf(source_path, output_txt_path)
            elif ext == ".docx":
                process_docx(source_path, output_txt_path)
            elif ext == ".txt":
                process_txt(source_path, output_txt_path)

            processed_count += 1
            if ocr_used_flag:
                ocr_count += 1

        except Exception as e_loop:
            print(f"\n[ERRO FATAL NO LOOP] Arquivo: {source_path} - Erro: {e_loop}")
            traceback.print_exc()
            error_count += 1

    print("-" * 30)
    print("Pré-processamento concluído!")
    print(f"Arquivos verificados: {len(files_to_process)}")
    print(f"Arquivos processados/sobrescritos: {processed_count}")
    print(f"Arquivos que necessitaram OCR: {ocr_count}")
    print(f"Erros encontrados: {error_count}")
    print(f"Resultados salvos em: {OUTPUT_DIR}")