# -*- coding: utf-8 -*-
import os
import shutil
import streamlit as st
# <<-- Imports Reativados -->>
import PyPDF2
from docx import Document
# <<-- Fim Imports Reativados -->>
import datetime
import pandas as pd

# --- Langchain Imports ---
import langchain # <--- ADICIONADO PARA RESOLVER O NameError
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.llms import OpenAI # Mantido para compatibilidade com text-davinci-003
from langchain.schema import Document as LangChainDocument
import chromadb
from chromadb import EphemeralClient
# --- OpenAI Imports ---
from openai import OpenAIError, OpenAI as OpenAIClient # Cliente principal renomeado

# --- Outros Imports ---
import sys
import glob
import traceback # Adicionado para melhor depuração
import tiktoken # Para cálculo de tokens
import re # Import já existente, necessário para preprocessamento
import unicodedata # Adicionado para normalização Unicode no preprocessamento
import chardet # Import necessário para carregar_txt

# Adicionar caminhos ao sys.path para garantir que os módulos sejam encontrados
try:
    # Assume __file__ exists when run as script
    script_dir = os.path.dirname(os.path.abspath(__file__))
    if script_dir not in sys.path:
        sys.path.append(script_dir)
except NameError:
    # Fallback for environments where __file__ is not defined (e.g., interactive)
    script_dir = os.getcwd()
    if script_dir not in sys.path:
         sys.path.append(script_dir)


# --- Configurações de caminhos ---
BASE_FONTES = r"C:\Users\pbm_s\OneDrive\Prova Digital" # Pasta raiz (AJUSTE CONFORME NECESSÁRIO)
OCR_OUTPUT_DIR = os.path.join(BASE_FONTES, "ocr_texts") # Pasta com os TXTs gerados pelo OCR
SUBPASTA_TEMP = os.path.join(BASE_FONTES, "uploads_temp")
PERSIST_DIR = os.path.join(BASE_FONTES, "chroma") # Banco de dados vetorial principal
HISTORICO_PATH = os.path.join(BASE_FONTES, "historico_respostas.txt")
FEEDBACK_PATH = os.path.join(BASE_FONTES, "feedback_respostas.csv")
FEEDBACK_FILES_DIR = os.path.join(BASE_FONTES, "feedback_files")
FEEDBACK_CONTENT_DIR = os.path.join(BASE_FONTES, "feedback_content")
EXEMPLOS_INTERNOS_DIR = os.path.join(BASE_FONTES, "exemplos_internos")


# Cria diretórios necessários (se não existirem)
os.makedirs(BASE_FONTES, exist_ok=True)
os.makedirs(OCR_OUTPUT_DIR, exist_ok=True)
os.makedirs(SUBPASTA_TEMP, exist_ok=True)
os.makedirs(PERSIST_DIR, exist_ok=True)
os.makedirs(FEEDBACK_FILES_DIR, exist_ok=True)
os.makedirs(FEEDBACK_CONTENT_DIR, exist_ok=True)
os.makedirs(EXEMPLOS_INTERNOS_DIR, exist_ok=True)

# ⚠️ Define a API Key fixa ANTES da verificação (RISCO DE SEGURANÇA) ⚠️
# Substitua pela sua chave real ou use variáveis de ambiente / segredos do Streamlit
# Exemplo usando segredos do Streamlit (recomendado):
# OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
# os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
# OU para teste local rápido (NÃO RECOMENDADO PARA PRODUÇÃO):
# os.environ["OPENAI_API_KEY"] = "sk-proj...-"

# --- Configuração e Verificação da Chave API via Streamlit Secrets ---
try:
    openai_api_key = st.secrets["OPENAI_API_KEY"]
    if not openai_api_key or openai_api_key == "sk-..." or "COLOQUE" in openai_api_key:
         st.error("Erro: A chave 'OPENAI_API_KEY' encontrada nos segredos está vazia ou parece ser um placeholder.")
         st.info("Verifique seu arquivo .streamlit/secrets.toml ou as configurações de segredos na Streamlit Cloud.")
         st.stop()
    os.environ["OPENAI_API_KEY"] = openai_api_key
    # st.sidebar.success("Chave API carregada dos segredos!") # Opcional
except KeyError:
    st.error("Erro Crítico: Chave 'OPENAI_API_KEY' não encontrada nos segredos do Streamlit.")
    st.info("Certifique-se de criar o arquivo .streamlit/secrets.toml e adicionar sua chave, ou configure os segredos na Streamlit Cloud.")
    st.stop()
except Exception as e:
    st.error(f"Erro inesperado ao acessar os segredos do Streamlit: {e}")
    st.stop()
# -------------------------------------------------------------------

# Inicializa cliente OpenAI e embeddings
# Renomeado cliente para evitar conflito com a classe OpenAI da langchain legada
openai_client_instance = None
embeddings = None
try:
    openai_client_instance = OpenAIClient() # Usa a chave da variável de ambiente por padrão
    embeddings = OpenAIEmbeddings()
except Exception as e:
    st.error(f"Erro ao inicializar cliente OpenAI ou embeddings: {str(e)}")
    st.stop()

# --- Definições de limites de CONTEXTO TOTAL dos modelos ---
def obter_limite_contexto(model_name):
    limites = {
        "text-davinci-003": 4097, "gpt-3.5-turbo": 16385, "gpt-4": 8192,
        "gpt-4-turbo": 128000, "gpt-4o": 128000,
        "gpt-4.1": 1048576, "gpt-4.1-mini": 1048576, "gpt-4.1-nano": 1048576, # Limites hipotéticos/exemplo
    }
    # Extrai o nome base caso venha com sufixos como -preview
    model_base = model_name.split('-')[0] + '-' + model_name.split('-')[1] if '-' in model_name else model_name
    default = 8192 if 'gpt-4' in model_base else (16385 if 'gpt-3.5' in model_base else 4097)
    return limites.get(model_name, limites.get(model_base, default))


# --- Definições de limites de OUTPUT dos modelos ---
def obter_limite_output(model_name):
    limites_output = {
        "gpt-4o": 4096, "gpt-4-turbo": 4096, "gpt-4": 8192,
        "gpt-3.5-turbo": 4096, "text-davinci-003": 4097,
        "gpt-4.1": 32768, "gpt-4.1-mini": 32768, "gpt-4.1-nano": 32768, # Limites hipotéticos/exemplo
    }
    model_base = model_name.split('-')[0] + '-' + model_name.split('-')[1] if '-' in model_name else model_name
    # Default mais conservador para GPTs (4096) a menos que especificado
    default_limit = 4096 if "gpt" in model_base else obter_limite_contexto(model_base)
    return limites_output.get(model_name, limites_output.get(model_base, default_limit))

# Função para estimar tokens
def estimar_tokens(texto, model_name="gpt-4o"):
    if not texto: return 0
    try:
        # Tenta obter o encoding específico do modelo
        encoding = tiktoken.encoding_for_model(model_name)
    except KeyError:
        # Fallback para um encoding comum se o modelo específico não for encontrado
        try:
            encoding = tiktoken.get_encoding("cl100k_base")
        except Exception:
            # Fallback final se tiktoken falhar completamente
            return len(texto) // 4 # Aproximação muito grosseira
    except Exception:
         return len(texto) // 4 # Outro fallback
    try:
        # Codifica o texto para obter os tokens
        return len(encoding.encode(texto))
    except Exception:
        # Fallback em caso de erro na codificação
        return len(texto) // 4

# Função para criar e executar LLM
def criar_e_executar_llm(prompt, model_name, temperature=0.2, max_tokens=1000):
    global openai_client_instance # Usa a instância renomeada
    if openai_client_instance is None:
        st.error("Cliente OpenAI (openai_client_instance) não inicializado.")
        raise Exception("OpenAI Client not initialized")

    # Garante que max_tokens seja pelo menos 1, se não for None
    if max_tokens is not None and max_tokens < 1:
        st.warning(f"max_tokens ajustado de {max_tokens} para 1.")
        max_tokens = 1
    # Se max_tokens for 0 ou None de forma inesperada, pode definir um mínimo padrão
    if max_tokens is None or max_tokens == 0:
         max_tokens = 100 # Ou outro valor padrão razoável

    try:
        # Verifica se é um modelo de Chat (GPT-3.5, GPT-4, etc.)
        if "gpt-4" in model_name or "gpt-3.5" in model_name or "gpt-4o" in model_name:
            messages = [
                {"role": "system", "content": "Você é um assistente especializado em análise jurídica e técnica de provas digitais."},
                {"role": "user", "content": prompt}
            ]
            response = openai_client_instance.chat.completions.create(
                model=model_name,
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens
            )
            # Verifica se a resposta e a mensagem existem
            if response.choices and response.choices[0].message and response.choices[0].message.content:
                return response.choices[0].message.content
            else:
                st.error(f"Resposta inesperada da API para {model_name}: {response}")
                return "[Erro: Resposta API inesperada ou vazia]"
        # Verifica se é um modelo de Completions legado (Davinci)
        elif "davinci" in model_name:
             st.warning(f"Usando modelo legado '{model_name}'. Considere migrar para modelos mais recentes.")
             try:
                 # Usa a API de Completions legada (via cliente OpenAI)
                 response = openai_client_instance.completions.create(
                    model=model_name,
                    prompt=prompt,
                    temperature=temperature,
                    max_tokens=max_tokens
                 )
                 if response.choices and response.choices[0].text:
                     return response.choices[0].text
                 else:
                     st.error(f"Resposta inesperada (legado) para {model_name}: {response}")
                     return "[Erro: Resposta API legada inesperada ou vazia]"
             except Exception as legacy_err:
                 st.error(f"Erro ao chamar modelo legado {model_name}: {legacy_err}")
                 raise legacy_err
        else:
            # Modelo não suportado
            raise ValueError(f"Modelo '{model_name}' não é suportado ou reconhecido.")

    except OpenAIError as api_err:
        st.error(f"Erro na API OpenAI ({model_name}): {api_err}")
        raise api_err # Re-levanta para ser tratado no bloco principal
    except Exception as e:
        st.error(f"Erro genérico ao executar LLM ({model_name}): {str(e)}")
        st.code(traceback.format_exc()) # Mostra mais detalhes do erro
        raise e # Re-levanta

# Função para testar a API OpenAI
def testar_openai(model_name="gpt-4o"):
    if not os.getenv("OPENAI_API_KEY"):
        return False, "Chave da API OpenAI não configurada."
    global openai_client_instance
    if openai_client_instance is None:
        return False, "Cliente OpenAI (openai_client_instance) não inicializado."
    try:
        # Usa um prompt simples e espera uma resposta curta
        resposta = criar_e_executar_llm(
            prompt="Teste de API. Responda apenas com a palavra 'OK'.",
            model_name=model_name,
            temperature=0,
            max_tokens=10
        )
        resposta_limpa = resposta.strip().lower() if resposta else ""
        # Verifica se a resposta contém 'ok' (mais flexível que igualdade exata)
        if "ok" in resposta_limpa:
            return True, f"Conexão OK com {model_name}. Resposta recebida: '{resposta.strip()}'"
        else:
            return False, f"Resposta inesperada de {model_name}: '{resposta.strip() if resposta else 'N/A'}'. Esperava conter 'OK'."
    except Exception as e:
        return False, f"Erro durante o teste de conexão com {model_name}: {str(e)}"

#======================================================================
# FUNÇÃO DE PRÉ-PROCESSAMENTO DE TEXTO
#======================================================================
def preprocessar_texto(texto: str) -> str:
    if not texto or not isinstance(texto, str):
        return ""
    try:
        # Remove quebras de linha que não separam parágrafos
        texto = re.sub(r'(?<!\n)\n(?!\n)', ' ', texto)
        # Remove hífens no final de linha (comuns em OCR)
        texto = re.sub(r'-\n', '', texto)
        # Normaliza espaços em branco múltiplos para um único espaço
        texto = re.sub(r'\s+', ' ', texto).strip()
        # Divide em linhas para remover linhas que contêm apenas números (ex: números de página)
        linhas = texto.split('\n')
        padrao_numero_linha = re.compile(r'^\s*\d+\s*$') # Linha com apenas número e espaços
        linhas_filtradas = [linha for linha in linhas if not padrao_numero_linha.fullmatch(linha.strip())]
        texto = '\n'.join(linhas_filtradas)
        # Normalização Unicode para consistência (ex: remove ligaturas)
        texto = unicodedata.normalize('NFKC', texto)
        # Garante que parágrafos sejam separados por duas quebras de linha, removendo excessos
        texto = re.sub(r'\n\s*\n', '\n\n', texto).strip()
    except Exception as e:
        # Log do erro para depuração, mas retorna o texto original (ou parcialmente processado)
        print(f"Erro durante o pré-processamento: {e}")
        # Considerar logar o erro em um arquivo ou sistema de log mais robusto
        # st.warning(f"Erro no pré-processamento: {e}") # Pode poluir a interface
        return texto # Retorna o texto como está se houver erro
    return texto
#======================================================================
# FIM PRÉ-PROCESSAMENTO
#======================================================================


# --- Funções de Carregamento (PDF, DOCX, TXT) ---

def carregar_pdf(path):
    """Carrega texto de PDF, lida com criptografia básica e aplica pré-processamento."""
    raw_text = ""
    metadata = {"source": os.path.abspath(path), "tipo": "documento_pdf"} # Usa path absoluto
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            # Verifica criptografia
            if reader.is_encrypted:
                try:
                    # Tenta descriptografar com senha vazia (comum)
                    reader.decrypt('')
                except Exception:
                    # Se falhar, marca como erro e retorna
                    st.warning(f"PDF '{os.path.basename(path)}' criptografado e senha incorreta/não fornecida.")
                    raw_text = f"[Erro: PDF criptografado: {os.path.basename(path)}]"
                    metadata["tipo"] = "erro_leitura_pdf_criptografado"
                    # Não precisa pré-processar a mensagem de erro
                    return [{"page_content": raw_text, "metadata": metadata}]

            # Extrai texto página por página
            raw_text_pages = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    raw_text_pages.append(page_text) # Adiciona texto da página
                else:
                    # Marca página sem texto extraível (pode ser imagem)
                    raw_text_pages.append(f"[Página {i+1} sem texto extraível ou vazia]")
            # Junta o texto das páginas com um separador claro
            raw_text = "\n\n--- Quebra de Página ---\n\n".join(raw_text_pages)

    except FileNotFoundError:
        st.warning(f"Erro: Arquivo PDF não encontrado em {path}")
        raw_text = f"[Erro: PDF não encontrado: {os.path.basename(path)}]"
        metadata["tipo"] = "erro_arquivo_nao_encontrado"
    except Exception as e:
        st.warning(f"Erro ao ler PDF {os.path.basename(path)}: {str(e)}")
        raw_text = f"[Erro ao processar PDF: {str(e)} - {traceback.format_exc(limit=1)}]"
        metadata["tipo"] = "erro_leitura_pdf"

    # Aplica pré-processamento apenas se não for uma mensagem de erro
    if not raw_text.startswith("[Erro"):
        processed_text = preprocessar_texto(raw_text)
    else:
        processed_text = raw_text # Mantém a mensagem de erro

    return [{"page_content": processed_text, "metadata": metadata}]


def carregar_docx(path):
    """Carrega texto de DOCX, incluindo tabelas básicas, e aplica pré-processamento."""
    raw_text = ""
    metadata = {"source": os.path.abspath(path), "tipo": "documento_docx"} # Usa path absoluto
    try:
        doc = Document(path) # Usa a classe Document importada
        # Extrai texto dos parágrafos
        full_text_list = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
        raw_text = "\n".join(full_text_list)

        # Tenta extrair texto de tabelas
        try:
             if doc.tables:
                 raw_text += "\n\n--- Início Bloco de Tabelas ---\n"
                 for table_idx, table in enumerate(doc.tables):
                      raw_text += f"\n--- Tabela {table_idx+1} ---\n"
                      header_done = False
                      for row_idx, row in enumerate(table.rows):
                           row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells] # Limpa células
                           row_text = "\t|\t".join(row_cells)
                           raw_text += row_text + "\n"
                           # Adiciona separador de cabeçalho (opcional, mas útil)
                           if not header_done and row_idx == 0 and len(row_cells) > 0:
                               raw_text += "--- | --- " * len(row_cells) + "\n"
                               header_done = True
                      raw_text += f"--- Fim Tabela {table_idx+1} ---\n"
                 raw_text += "\n--- Fim Bloco de Tabelas ---\n"
        except Exception as e_table:
            # Apenas informa, não impede o processamento do resto do texto
            st.info(f"Info: Não foi possível extrair tabelas de {os.path.basename(path)} (ou erro na formatação): {e_table}")

    except FileNotFoundError:
        st.warning(f"Erro: Arquivo DOCX não encontrado em {path}")
        raw_text = f"[Erro: DOCX não encontrado: {os.path.basename(path)}]"
        metadata["tipo"] = "erro_arquivo_nao_encontrado"
    except Exception as e:
        st.warning(f"Erro ao ler DOCX {os.path.basename(path)}: {str(e)}")
        raw_text = f"[Erro ao processar DOCX: {str(e)} - {traceback.format_exc(limit=1)}]"
        metadata["tipo"] = "erro_leitura_docx"

    # Aplica pré-processamento apenas se não for uma mensagem de erro
    if not raw_text.startswith("[Erro"):
        processed_text = preprocessar_texto(raw_text)
    else:
        processed_text = raw_text # Mantém a mensagem de erro

    return [{"page_content": processed_text, "metadata": metadata}]


def carregar_txt(path):
    """Carrega um arquivo TXT, detecta codificação e aplica pré-processamento."""
    raw_text = ""
    metadata = {"source": os.path.abspath(path), "tipo": "documento_texto"} # Usa path absoluto
    try:
        # Detecta a codificação lendo uma parte inicial do arquivo binário
        with open(path, "rb") as f_raw:
             # Ler mais bytes pode ajudar na detecção, mas aumenta I/O
             raw_bytes = f_raw.read(40960) # Lê até 40KB para detecção
             detection = chardet.detect(raw_bytes)
             detected_encoding = detection['encoding']
             confidence = detection['confidence']
             # Usa UTF-8 como fallback se a detecção falhar ou tiver baixa confiança
             encoding_to_use = detected_encoding if detected_encoding and confidence > 0.7 else 'utf-8'
             # st.info(f"Lendo '{os.path.basename(path)}': Detectado {detected_encoding} (Conf: {confidence:.2f}), Usando: {encoding_to_use}")

        # Tenta ler o arquivo com a codificação detectada/fallback
        try:
            with open(path, "r", encoding=encoding_to_use, errors='ignore') as f: # errors='ignore' para evitar falha total
                raw_text = f.read()
        except (LookupError, UnicodeDecodeError) as e_decode:
             # Se a primeira tentativa falhar (mesmo com ignore), tenta UTF-8 explicitamente
             st.warning(f"Falha ao ler {os.path.basename(path)} com '{encoding_to_use}' ({e_decode}). Tentando utf-8.")
             try:
                 with open(path, "r", encoding='utf-8', errors='ignore') as f:
                     raw_text = f.read()
             except Exception as e_utf8:
                 st.error(f"Erro ao ler {os.path.basename(path)} mesmo com utf-8: {e_utf8}")
                 raw_text = f"[Erro: Falha na leitura de TXT com todas as codificações: {os.path.basename(path)}]"
                 metadata["tipo"] = "erro_leitura_txt_encoding"

    except FileNotFoundError:
         st.warning(f"Erro: Arquivo TXT não encontrado em {path}")
         raw_text = f"[Erro: Arquivo TXT não encontrado: {os.path.basename(path)}]"
         metadata["tipo"] = "erro_arquivo_nao_encontrado"
    except Exception as e:
        st.warning(f"Erro geral ao ler TXT {os.path.basename(path)}: {str(e)}")
        raw_text = f"[Erro ao processar TXT: {str(e)} - {traceback.format_exc(limit=1)}]"
        metadata["tipo"] = "erro_leitura_txt"

    # Aplica pré-processamento apenas se não for uma mensagem de erro
    if not raw_text.startswith("[Erro"):
        processed_text = preprocessar_texto(raw_text)
        # Verifica se o texto processado não ficou vazio
        if not processed_text:
             st.info(f"Arquivo TXT '{os.path.basename(path)}' resultou em texto vazio após pré-processamento.")
             # Pode retornar um marcador ou o texto original não processado
             processed_text = "[Conteúdo vazio após pré-processamento]"
             metadata["tipo"] = "aviso_texto_vazio_pos_processamento"
    else:
        processed_text = raw_text # Mantém a mensagem de erro

    return [{"page_content": processed_text, "metadata": metadata}]


# Função wrapper para carregar diferentes tipos de arquivo
def carregar_arquivo(path, is_attachment=False):
    st.warning(f"DEBUG (carregar_arquivo): Função chamada com path='{path}', is_attachment={is_attachment}") # <-- Adicionar log 1
    
    """
    Carrega um arquivo (TXT, PDF, DOCX).
    Para indexação principal (não anexo), geralmente se espera TXT da pasta OCR.
    Para anexos (consulta temporária ou feedback), tenta carregar PDF/DOCX/TXT.
    """
    # Validação básica do caminho
    if not path or not isinstance(path, str) or not os.path.exists(path) or not os.path.isfile(path):
         st.error(f"Caminho de arquivo inválido ou não encontrado: {path}")
         return [{"page_content": f"[Erro: Caminho inválido/não encontrado: {path}]",
                  "metadata": {"source": str(path), "tipo": "erro_caminho_invalido"}}]

    ext = os.path.splitext(path)[1].lower()
    nome_base = os.path.basename(path)
    st.warning(f"DEBUG (carregar_arquivo): Ext extraída='{ext}', nome_base='{nome_base}'") # <-- Adicionar log 2

    try:
        if ext == ".txt":
            return carregar_txt(path)
        elif ext == ".pdf":
            # Permite PDF se for anexo ou se explicitamente permitido para indexação principal
            if is_attachment: #  or ALLOW_PDF_INDEXING (se quisesse permitir)
                return carregar_pdf(path)
            else:
                 st.warning(f"PDF não é esperado para indexação principal (apenas de OCR): {nome_base}. Pulando.")
                 return [{"page_content": f"[Aviso: PDF ignorado para indexação principal: {nome_base}]", "metadata": {"source": path, "tipo": "formato_ignorado_principal"}}]
        elif ext == ".docx":
             # Permite DOCX se for anexo
            if is_attachment:
                return carregar_docx(path)
            else:
                 st.warning(f"DOCX não é esperado para indexação principal: {nome_base}. Pulando.")
                 return [{"page_content": f"[Aviso: DOCX ignorado para indexação principal: {nome_base}]", "metadata": {"source": path, "tipo": "formato_ignorado_principal"}}]
        else:
            # Formato não suportado
            st.warning(f"Formato não suportado: {ext} para o arquivo {nome_base}")
            erro_msg = f"[Erro: Formato não suportado: {ext}]"
            # Não precisa pré-processar mensagem de erro
            return [{"page_content": erro_msg, "metadata": {"source": path, "tipo": "formato_nao_suportado"}}]
    except Exception as e_load:
        st.error(f"Erro inesperado ao tentar carregar o arquivo {nome_base}: {e_load}")
        return [{"page_content": f"[Erro geral no carregamento: {e_load}]", "metadata": {"source": path, "tipo": "erro_carregamento_geral"}}]


# Função de divisão de documentos
def dividir_documentos(docs_com_texto_processado, chunk_size=512, chunk_overlap=100):
    """Divide documentos (já carregados e pré-processados) em chunks."""
    if not docs_com_texto_processado:
        st.info("Nenhum documento fornecido para divisão.")
        return []

    langchain_docs = []
    for doc_info in docs_com_texto_processado:
        page_content = ""
        metadata = {}

        # Extrai conteúdo e metadados, tratando dict ou objeto LangChainDocument
        if isinstance(doc_info, dict):
            page_content = doc_info.get("page_content", "")
            metadata = doc_info.get("metadata", {})
            # Garante que metadados sejam um dict
            if not isinstance(metadata, dict):
                metadata = {"source": str(metadata)} # Converte para dict se for só string
        elif isinstance(doc_info, LangChainDocument):
             page_content = doc_info.page_content
             metadata = doc_info.metadata
             if not isinstance(metadata, dict):
                 metadata = {"source": str(metadata)} # Converte para dict se for só string
        else:
            st.warning(f"Item inesperado recebido durante a divisão: {type(doc_info)}. Pulando.")
            continue # Pula este item

        # Adiciona à lista apenas se tiver conteúdo válido (não vazio e não erro/aviso)
        page_content_strip = page_content.strip() if page_content else ""
        if page_content_strip and not page_content_strip.startswith(("[Erro", "[Aviso", "[Conteúdo vazio")):
             # Cria o objeto LangChainDocument
             langchain_docs.append(LangChainDocument(page_content=page_content, metadata=metadata))
        # else:
            # Opcional: Logar por que um documento foi pulado
            # if page_content_strip.startswith("["):
            #     st.info(f"Documento ignorado na divisão (marcador/erro): {metadata.get('source', 'origem desconhecida')}")
            # elif not page_content_strip:
            #      st.info(f"Documento ignorado na divisão (conteúdo vazio): {metadata.get('source', 'origem desconhecida')}")

    if not langchain_docs:
        st.info("Nenhum conteúdo válido encontrado nos documentos para divisão.")
        return []

    try:
        # Configura o divisor de texto
        # RecursiveCharacterTextSplitter é geralmente bom para textos gerais
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,         # Tamanho alvo do chunk
            chunk_overlap=chunk_overlap,   # Sobreposição entre chunks
            length_function=len,           # Função para medir o tamanho (padrão: len)
            add_start_index=True,          # Adiciona índice inicial do chunk no metadado (útil)
            separators=["\n\n", "\n", ". ", ", ", " ", ""] # Separadores priorizados
            )

        # st.info(f"Dividindo {len(langchain_docs)} documento(s) em chunks (size={chunk_size}, overlap={chunk_overlap})...")
        docs_divididos = splitter.split_documents(langchain_docs)
        st.info(f"Divisão concluída: {len(langchain_docs)} documento(s) -> {len(docs_divididos)} chunks.")
        return docs_divididos
    except Exception as e:
        st.error(f"Erro crítico ao dividir documentos: {e}")
        st.code(traceback.format_exc())
        return [] # Retorna lista vazia em caso de erro

# Função de carregar ChromaDB
def carregar_chroma():
    """Carrega ou inicializa a instância do ChromaDB."""
    global embeddings
    if embeddings is None:
        st.error("Embeddings não inicializados. Não é possível carregar ChromaDB.")
        return None

    persist_directory = PERSIST_DIR # Usa a variável global

    # Verifica se o diretório existe
    if not os.path.exists(persist_directory):
        st.info(f"Diretório ChromaDB não encontrado em '{persist_directory}'. Tentando criar...")
        try:
            os.makedirs(persist_directory, exist_ok=True)
            st.success(f"Diretório ChromaDB criado: {persist_directory}")
            # Tenta inicializar uma nova instância vazia
            st.info("Inicializando nova instância ChromaDB vazia...")
            # O próprio Chroma lida com a criação se o diretório estiver vazio
            db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
            # Adiciona e remove um documento dummy para garantir que a coleção seja criada
            dummy_id = f"dummy_{datetime.datetime.now().timestamp()}"
            db.add_texts(["init"], metadatas=[{"source":"dummy"}], ids=[dummy_id])
            db.delete(ids=[dummy_id])
            db.persist() # Importante persistir após modificações
            st.success("Nova instância ChromaDB inicializada com sucesso.")
            return db
        except Exception as create_err:
            st.error(f"Falha crítica ao criar ou inicializar ChromaDB: {create_err}")
            st.code(traceback.format_exc())
            return None
    else:
        # Se o diretório existe, tenta carregar a instância existente
        try:
            # st.info(f"Carregando ChromaDB de: {persist_directory}")
            db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
            # Tenta uma operação simples para verificar se está funcionando
            count = db._collection.count()
            # st.success(f"ChromaDB carregado com sucesso ({count} itens).")
            return db
        except Exception as load_err:
            # Erros comuns incluem coleção não encontrada ou banco corrompido
            error_str = str(load_err)
            # Verifica mensagens de erro comuns de DB vazio ou não inicializado
            is_empty_error = ("Could not find collection" in error_str or
                              "zero items" in error_str or
                              "does not exist" in error_str or
                              "no such table" in error_str) # Adiciona verificação SQLite
            if is_empty_error:
                st.warning(f"Base ChromaDB parece vazia ou inválida ({error_str}). Pode ser necessário reindexar ou reparar.")
                # Tenta retornar uma nova instância que pode criar a coleção sob demanda
                try:
                    return Chroma(persist_directory=persist_directory, embedding_function=embeddings)
                except Exception as e_recreate:
                     st.error(f"Falha ao tentar recriar instância Chroma após erro de vazio: {e_recreate}")
                     return None
            else:
                st.error(f"Erro crítico ao carregar ChromaDB existente: {load_err}")
                st.code(traceback.format_exc())
                st.info("Considere usar a opção 'Reparar ChromaDB' na barra lateral se o problema persistir.")
                return None

# Função de diagnóstico do ChromaDB
def diagnosticar_chromadb():
    st.subheader("Diagnóstico do ChromaDB")
    global embeddings
    if embeddings is None:
        st.error("❌ Embeddings não inicializados. Diagnóstico abortado.")
        return

    persist_directory = PERSIST_DIR
    st.write(f"**Diretório Persistente:** `{persist_directory}`")

    # 1. Verificar existência do diretório
    if not os.path.exists(persist_directory):
        st.error(f"❌ Diretório ChromaDB não existe: {persist_directory}")
        return
    else:
        st.success(f"✅ Diretório ChromaDB encontrado.")
        try:
            files = os.listdir(persist_directory)
            st.info(f"Conteúdo do diretório: {len(files)} itens.")
            # Opcional: listar alguns arquivos (pode ser útil para depuração)
            # st.text('\n'.join(files[:5]) + ('...' if len(files) > 5 else ''))
        except Exception as e:
            st.error(f"❌ Erro ao listar conteúdo do diretório: {e}")

    # 2. Verificar versões das bibliotecas
    st.markdown("**Versões das Bibliotecas:**")
    try:
        # Importa explicitamente aqui para garantir que temos os módulos corretos
        import langchain as lchain_diag
        import chromadb as cdb_diag
        import openai as openai_diag
        st.text(f"- LangChain: `{lchain_diag.__version__}`")
        st.text(f"- ChromaDB: `{cdb_diag.__version__}`")
        st.text(f"- OpenAI: `{openai_diag.__version__}`")
    except Exception as e_ver:
        st.warning(f"Não foi possível verificar todas as versões das bibliotecas: {e_ver}")

    # 3. Tentar carregar e obter informações
    st.info("Tentando carregar instância e contar itens...")
    db_instance_diag = None
    try:
        db_instance_diag = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
        # Verifica se a coleção existe antes de contar
        # Nota: Acessar _collection diretamente é um detalhe de implementação e pode mudar.
        # Uma abordagem mais segura seria usar um método público se disponível, ou tratar a exceção.
        try:
             count = db_instance_diag._collection.count()
             st.success(f"✅ ChromaDB carregado com sucesso! Itens indexados: **{count}**")
        except Exception as e_count:
             # Verifica se o erro é por coleção inexistente
             if "does not exist" in str(e_count) or "Could not find collection" in str(e_count):
                  st.warning("Coleção padrão do ChromaDB não encontrada. A base está vazia ou precisa ser inicializada/reparada.")
                  count = 0 # Assume 0 itens se a coleção não existe
             else:
                  raise e_count # Re-levanta outros erros de contagem

        # 4. Tentar obter um exemplo de metadado (se houver itens)
        if count > 0:
             try:
                 st.info("Obtendo exemplo de metadado (1 item)...")
                 # include=['metadatas', 'documents'] para ver mais detalhes
                 info = db_instance_diag.get(limit=1, include=['metadatas'])
                 if info and info.get('metadatas') and len(info['metadatas']) > 0:
                     st.text("Exemplo de Metadado:")
                     st.json(info['metadatas'][0])
                 else:
                      st.warning("Não foi possível obter exemplo de metadado (retorno vazio ou inválido).")
             except Exception as get_err:
                 st.warning(f"Não foi possível obter exemplo de metadado: {get_err}")
        elif count == 0: # Condição explícita para 0 itens
            st.info("Base de dados está vazia (0 itens).")

    except Exception as e_load_diag:
        st.error(f"❌ Erro crítico ao carregar ou consultar ChromaDB durante diagnóstico: {str(e_load_diag)}")
        st.code(traceback.format_exc())
        st.info("Isso pode indicar um banco de dados corrompido, um problema de configuração ou versão. Considere usar a opção 'Reparar'.")


# Função de indexação de documentos já divididos
def indexar_documentos(docs_chunked):
    """Indexa uma lista de chunks (LangChainDocument) no ChromaDB."""
    if not docs_chunked:
        st.info("Nenhum chunk fornecido para indexação.")
        return 0 # Retorna 0 chunks indexados

    global embeddings
    if embeddings is None:
        st.error("Embeddings não inicializados. Indexação abortada.")
        return 0

    # Garante que temos uma lista de LangChainDocuments
    if not isinstance(docs_chunked, list) or not all(isinstance(doc, LangChainDocument) for doc in docs_chunked):
        st.error("Erro: a função indexar_documentos espera uma lista de objetos LangChainDocument.")
        return 0

    try:
        # Carrega a instância do ChromaDB
        db = carregar_chroma()
        if db is None:
            st.error("Falha ao carregar/inicializar ChromaDB. Indexação abortada.")
            return 0

        total_chunks_to_process = len(docs_chunked)
        st.info(f"Iniciando indexação de {total_chunks_to_process} chunks...")

        # Barra de progresso e status
        progress_bar_index = st.progress(0.0)
        status_text_index = st.empty()
        status_text_index.text(f"Preparando indexação de 0/{total_chunks_to_process} chunks...")

        # Indexação em lotes (batch) para eficiência e evitar timeouts
        batch_size = 100 # Ajustável: lotes menores consomem menos memória, maiores podem ser mais rápidos
        total_chunks_indexed = 0
        erros_indexacao = 0

        for i in range(0, total_chunks_to_process, batch_size):
            current_batch = docs_chunked[i:min(i + batch_size, total_chunks_to_process)]
            if not current_batch: continue # Lote vazio (não deve acontecer, mas por segurança)

            # Gera IDs únicos para cada chunk no lote
            # Usar hash do conteúdo + source pode ajudar na deduplicação, mas IDs únicos são mais simples
            batch_ids = []
            batch_texts = []
            batch_metadatas = []
            for j, doc in enumerate(current_batch):
                 # Garante que metadados existam e tenham 'source'
                 doc_meta = doc.metadata if isinstance(doc.metadata, dict) else {}
                 if 'source' not in doc_meta or not doc_meta['source']:
                      doc_meta['source'] = 'fonte_desconhecida_' + str(i+j) # Adiciona ID se faltar source

                 source_name = os.path.basename(str(doc_meta.get('source', '')))
                 # Cria ID mais robusto: source + hash do conteúdo (primeiros N chars) + timestamp
                 content_prefix_hash = hash(doc.page_content[:100])
                 unique_id = f"src_{source_name[:20]}_h_{content_prefix_hash}_{i+j}_{datetime.datetime.now().timestamp():.6f}"
                 # Limpa caracteres inválidos para IDs do Chroma (ex: espaços, /, ?)
                 safe_id = re.sub(r'[^\w\._-]+', '_', unique_id) # Permite alfanuméricos, _, ., -
                 safe_id = safe_id[:100] # Limita tamanho do ID se necessário

                 batch_ids.append(safe_id)
                 batch_texts.append(doc.page_content)
                 batch_metadatas.append(doc_meta)


            try:
                 # Adiciona o lote ao ChromaDB usando add_texts para melhor controle
                 db.add_texts(texts=batch_texts, metadatas=batch_metadatas, ids=batch_ids)

                 total_chunks_indexed += len(current_batch)
                 progress_percentage = min(1.0, total_chunks_indexed / total_chunks_to_process)
                 progress_bar_index.progress(progress_percentage)
                 status_text_index.text(f"Indexando {total_chunks_indexed}/{total_chunks_to_process} chunks...")

            except Exception as batch_error:
                 erros_indexacao += 1
                 st.warning(f"Erro ao indexar lote iniciado em {i} (ID exemplo: {batch_ids[0] if batch_ids else 'N/A'}): {batch_error}")
                 # Opcional: Logar mais detalhes do erro ou dos documentos problemáticos
                 # st.text(f"Primeiro texto do lote com erro:\n{batch_texts[0][:200]}...")
                 continue # Tenta o próximo lote

        # Limpa barra de progresso e texto de status
        progress_bar_index.empty(); status_text_index.empty()

        # Persiste as alterações no disco APÓS processar todos os lotes
        if total_chunks_indexed > 0:
             try:
                 st.info("Persistindo alterações no ChromaDB...")
                 db.persist()
                 st.success(f"Indexação concluída e persistida! {total_chunks_indexed} chunks adicionados/atualizados.")
                 if erros_indexacao > 0:
                     st.warning(f"{erros_indexacao} lote(s) encontraram erros durante a indexação.")
             except Exception as persist_error:
                 st.error(f"Erro crítico ao persistir ChromaDB após indexação: {persist_error}")
                 st.code(traceback.format_exc())
                 # Mesmo com erro de persist, os dados podem estar na memória
        elif erros_indexacao == 0:
             st.info("Nenhum chunk novo foi processado para indexação (ou todos os lotes falharam sem chunks válidos).")
        else: # Apenas erros
             st.error(f"Indexação falhou. {erros_indexacao} lote(s) encontraram erros.")


        return total_chunks_indexed

    except Exception as e:
        st.error(f"Erro geral durante o processo de indexação: {str(e)}")
        st.code(traceback.format_exc())
        # Garante limpeza da UI em caso de erro
        if 'progress_bar_index' in locals(): progress_bar_index.empty()
        if 'status_text_index' in locals(): status_text_index.empty()
        return 0


# Função principal de indexação da pasta
def indexar_pasta_completa(caminho_textos, recursive=True):
    """
    Varre a pasta `caminho_textos` por arquivos suportados (TXT),
    carrega, pré-processa, extrai metadados do nome (se possível),
    divide em chunks e chama a indexação final.
    """
    st.info(f"🔍 Iniciando varredura e indexação da pasta: {caminho_textos} (Recursivo: {recursive})")
    # Por enquanto, focado apenas em .txt vindos do OCR
    tipos_suportados = ["*.txt"]
    arquivos_encontrados = []

    # Valida o caminho da pasta
    if not os.path.isdir(caminho_textos):
         st.error(f"Erro Crítico: Diretório de textos '{caminho_textos}' não encontrado ou inválido.")
         return 0, 0 # Retorna 0 arquivos processados, 0 chunks indexados

    # Busca arquivos .txt
    for tipo in tipos_suportados:
         pattern = os.path.join(caminho_textos, "**", tipo) if recursive else os.path.join(caminho_textos, tipo)
         try:
             # Usa iglob para potencialmente economizar memória em pastas muito grandes
             arquivos_encontrados.extend(glob.iglob(pattern, recursive=recursive))
         except Exception as e_glob:
              st.error(f"Erro ao buscar arquivos com padrão '{pattern}': {e_glob}")
              return 0, 0

    # Filtra para garantir que são apenas arquivos (glob com ** pode retornar pastas)
    # Efetua a filtragem após coletar todos os caminhos
    arquivos_filtrados = [f for f in arquivos_encontrados if os.path.isfile(f)]

    if not arquivos_filtrados:
        st.warning(f"⚠️ Nenhum arquivo {tipos_suportados[0]} encontrado em '{caminho_textos}' (ou subpastas, se recursivo). Nada a indexar.")
        return 0, 0

    st.info(f"📚 {len(arquivos_filtrados)} arquivo(s) .txt encontrados para processar.")

    todos_os_chunks = [] # Lista para acumular chunks de todos os arquivos válidos
    arquivos_processados_ok = 0
    arquivos_com_erro = 0
    total_arquivos = len(arquivos_filtrados)

    # UI de Progresso para Leitura/Divisão
    progress_leitura = st.progress(0.0)
    status_text_leitura = st.empty()
    status_text_leitura.text(f"Lendo e processando arquivos .txt (0/{total_arquivos})...")

    # Parâmetros de chunking (podem ser configuráveis na UI no futuro)
    # Valores maiores de chunk_size podem ser melhores para RAG, dependendo do modelo
    chunk_config = {"chunk_size": 1024, "chunk_overlap": 150}
    st.info(f"Configuração de Chunking: Size={chunk_config['chunk_size']}, Overlap={chunk_config['chunk_overlap']}")

    for i, arquivo in enumerate(arquivos_filtrados):
        nome_arq = os.path.basename(arquivo)
        status_text_leitura.text(f"Processando {i+1}/{total_arquivos}: {nome_arq}")
        # st.text(f"Processando: {arquivo}") # Mais detalhe (opcional)

        try:
            # 1. Carrega o TXT (já aplica pré-processamento básico e detecção de encoding)
            # Retorna uma lista, geralmente com um dict contendo page_content e metadata
            docs_carregados_info = carregar_txt(arquivo)

            # Verifica se o carregamento foi bem-sucedido e retornou conteúdo válido
            conteudo_valido = False
            if (docs_carregados_info and isinstance(docs_carregados_info, list) and
                    len(docs_carregados_info) > 0 and
                    isinstance(docs_carregados_info[0], dict) and
                    docs_carregados_info[0].get("page_content") and
                    not docs_carregados_info[0]["page_content"].strip().startswith(("[Erro", "[Aviso", "[Conteúdo vazio"))):
                conteudo_valido = True

            if not conteudo_valido:
                st.warning(f"Leitura inválida, vazia ou erro para: {nome_arq}. Pulando.")
                arquivos_com_erro += 1
                continue # Pula para o próximo arquivo

            # 2. Extração Opcional de Metadados do Nome do Arquivo
            try:
                # Assume convenção: TIPO_ID_DETALHE*.txt
                # Ajuste o regex/split conforme a convenção real dos nomes de arquivo
                nome_sem_ext = os.path.splitext(nome_arq)[0]
                partes = nome_sem_ext.split('_') # Exemplo simples de split por underscore
                # Extração mais robusta pode usar regex
                tipo_doc_extraido = partes[0].upper() if len(partes) > 0 else 'DESCONHECIDO'
                id_doc_extraido = partes[1] if len(partes) > 1 else 'SEM_ID'
                # Adiciona metadados extraídos a CADA documento/chunk deste arquivo
                timestamp_idx = datetime.datetime.now().isoformat() # Timestamp único para este lote
                for doc_info in docs_carregados_info:
                    # Garante que 'metadata' existe e é um dict
                    if not isinstance(doc_info.get('metadata'), dict):
                        doc_info['metadata'] = {"source": arquivo} # Cria se não existir
                    doc_info['metadata']['tipo_documento_ext'] = tipo_doc_extraido
                    doc_info['metadata']['id_documento_ext'] = id_doc_extraido
                    # Adiciona timestamp da indexação
                    doc_info['metadata']['indexed_at'] = timestamp_idx

            except Exception as e_parse:
                st.info(f"Não foi possível extrair metadados do nome do arquivo: {nome_arq} ({e_parse}). Metadados básicos serão usados.")
                # Garante metadados mínimos se a extração falhar
                timestamp_idx = datetime.datetime.now().isoformat()
                for doc_info in docs_carregados_info:
                     if not isinstance(doc_info.get('metadata'), dict):
                          doc_info['metadata'] = {"source": arquivo}
                     doc_info['metadata'].setdefault('tipo_documento_ext', 'ERRO_PARSE') # Adiciona se não existir
                     doc_info['metadata'].setdefault('id_documento_ext', 'ERRO_PARSE')
                     doc_info['metadata']['indexed_at'] = timestamp_idx


            # 3. Divide o conteúdo carregado em Chunks usando a configuração definida
            # Passa a lista de dicts retornada por carregar_txt para dividir_documentos
            chunks_do_arquivo = dividir_documentos(docs_carregados_info, **chunk_config)

            # 4. Acumula os chunks gerados
            if chunks_do_arquivo:
                todos_os_chunks.extend(chunks_do_arquivo)
                arquivos_processados_ok += 1
            else:
                st.warning(f"Nenhum chunk válido gerado para {nome_arq} após divisão (conteúdo pode ser muito curto ou houve erro).")
                # Não conta como erro fatal, mas também não como sucesso
                # arquivos_com_erro += 1 # Descomentar se quiser contar isso como erro

        except Exception as e_proc_arq:
            st.error(f"Erro crítico inesperado ao processar o arquivo {nome_arq}: {str(e_proc_arq)}")
            st.code(traceback.format_exc(limit=1))
            arquivos_com_erro += 1

        # Atualiza progresso da leitura/divisão
        progress_leitura.progress(min(1.0, (i + 1) / total_arquivos))

    # Limpa UI de progresso da leitura
    progress_leitura.empty(); status_text_leitura.empty()

    # Verifica se algum chunk foi gerado no total
    if not todos_os_chunks:
        st.error(f"❌ Processamento concluído, mas nenhum chunk válido foi gerado a partir dos {total_arquivos} arquivos encontrados.")
        if arquivos_com_erro == total_arquivos:
             st.info("Todos os arquivos encontrados resultaram em erro ou conteúdo inválido/vazio.")
        return 0, 0

    st.success(f"Leitura e divisão concluídas: {arquivos_processados_ok}/{total_arquivos} arquivos válidos geraram {len(todos_os_chunks)} chunks.")
    if arquivos_com_erro > 0:
        st.warning(f"{arquivos_com_erro} arquivo(s) encontraram erros ou foram pulados.")

    st.info("Iniciando a indexação final dos chunks no ChromaDB...")

    # 5. Chama a função de indexação com TODOS os chunks acumulados
    chunks_indexados = indexar_documentos(todos_os_chunks)

    # Retorna o número de arquivos que geraram chunks e o total de chunks indexados
    return arquivos_processados_ok, chunks_indexados


# Função de listar arquivos indexados (baseado nos metadados)
def listar_arquivos_indexados():
    """Lista os nomes base dos arquivos fontes encontrados nos metadados do ChromaDB."""
    st.info("Listando fontes indexadas (baseado em amostra de metadados)...")
    try:
        db = carregar_chroma()
        if db is None:
            st.error("ChromaDB não carregado. Não é possível listar fontes.")
            return []

        count = 0
        try:
            count = db._collection.count()
        except Exception as e_count_list:
             st.warning(f"Não foi possível contar itens no DB para listar fontes ({e_count_list}). A base pode estar vazia ou corrompida.")
             return [] # Retorna vazio se não puder contar

        if count == 0:
            st.info("Base ChromaDB está vazia. Nenhuma fonte para listar.")
            return []

        # Obtém uma amostra maior se a base for grande, mas limita para performance
        limit_get = min(count, 2000)
        st.info(f"Obtendo metadados de até {limit_get} itens...")
        # Aumenta o timeout se necessário para bases grandes
        # import chromadb.config
        # settings = chromadb.config.Settings(...) pode ser usado se o client for criado manualmente
        # Mas com a classe Chroma, o controle é mais limitado.
        retrieved_data = db.get(limit=limit_get, include=['metadatas'])

        if not retrieved_data or not retrieved_data.get('metadatas'):
            st.warning("Nenhum metadado encontrado na amostra obtida.")
            return []

        # Processa os metadados para extrair fontes únicas
        sources = set()
        # Normaliza caminhos de diretórios internos para exclusão mais robusta
        internal_paths_norm = [os.path.normpath(p) for p in [EXEMPLOS_INTERNOS_DIR, FEEDBACK_CONTENT_DIR, FEEDBACK_FILES_DIR, SUBPASTA_TEMP] if p and os.path.exists(p)]

        for metadata in retrieved_data.get('metadatas', []):
            if isinstance(metadata, dict) and 'source' in metadata and isinstance(metadata['source'], str):
                 source_path = metadata['source']
                 # Verifica se o caminho é válido antes de normalizar
                 if not source_path: continue

                 try:
                    norm_source_path = os.path.normpath(source_path)
                 except Exception:
                     # st.warning(f"Metadado 'source' inválido encontrado durante listagem: {source_path}") # Muito verbose
                     continue # Pula fonte inválida

                 # Verifica se a fonte pertence a um diretório interno (feedback, exemplo, temp)
                 is_internal = False
                 if "dummy" in norm_source_path or "init_marker" in norm_source_path:
                     is_internal = True
                 else:
                    try:
                        # Verifica se o caminho da fonte começa com algum dos caminhos internos normalizados
                        is_internal = any(norm_source_path.startswith(ipn + os.sep) for ipn in internal_paths_norm)
                    except Exception as e_path_cmp:
                        # Erro na comparação de caminhos (raro, mas possível com caracteres estranhos)
                         # st.warning(f"Erro ao comparar caminho da fonte '{norm_source_path}' com internos: {e_path_cmp}")
                         continue # Pula em caso de erro na verificação

                 # Adiciona à lista se não for interno e se o arquivo ainda existir (opcional)
                 # check_exists = True # Descomentar para verificar se o arquivo original ainda existe
                 # if not is_internal and (not check_exists or os.path.exists(norm_source_path)):
                 if not is_internal:
                     sources.add(norm_source_path) # Adiciona o caminho completo normalizado

        # Retorna apenas o nome base dos arquivos ordenados
        return sorted([os.path.basename(f) for f in sources if f]) # Garante que f não seja vazio

    except Exception as e:
        st.error(f"Erro inesperado ao listar fontes indexadas: {str(e)}")
        st.code(traceback.format_exc())
        return []

# Função de combinar resultados RAG
def combinar_resultados(retrievers, query, k=7, incluir_feedback=True):
    """Busca documentos relevantes de múltiplos retrievers e combina os resultados."""
    if not retrievers:
        st.warning("Nenhum retriever fornecido para busca RAG.")
        return []

    all_retrieved_docs = []
    docs_feedback = []

    # 1. Busca em todos os retrievers fornecidos (principal e temporário, se houver)
    st.info(f"Buscando em {len(retrievers)} retriever(s) com k={k}...")
    for i, retriever in enumerate(retrievers):
        retriever_name = f"Retriever {i+1}" # Pode melhorar identificando (Principal/Temp)
        try:
            # get_relevant_documents é a função padrão do Langchain Retriever
            retrieved_docs = retriever.get_relevant_documents(query)
            st.info(f"{retriever_name} retornou {len(retrieved_docs)} documentos.")
            # Adiciona metadado para saber de qual retriever veio (útil para debug/ranking)
            for doc in retrieved_docs:
                if not hasattr(doc, 'metadata') or not isinstance(doc.metadata, dict):
                    doc.metadata = {} # Garante que metadados existam como dict
                doc.metadata['_retriever_index'] = i
                doc.metadata['_retriever_name'] = retriever_name
            all_retrieved_docs.extend(retrieved_docs)
        except Exception as e:
            st.warning(f"Erro ao buscar no {retriever_name}: {str(e)}")
            # Pode adicionar mais detalhes do erro se necessário

    # 2. Busca Específica por Feedback (se habilitado e se houver retriever principal)
    if incluir_feedback and retrievers:
        try:
            # Assume que o primeiro retriever é o principal (ChromaDB persistente)
            if hasattr(retrievers[0], 'vectorstore') and retrievers[0].vectorstore is not None:
                db_vectorstore = retrievers[0].vectorstore
                st.info("Buscando documentos de feedback adicionais...")
                # Tenta buscar com filtro específico para tipo 'feedback' ou 'feedback_attachment'
                try:
                    # Filtro mais robusto para ChromaDB (pode precisar de ajuste dependendo da versão)
                    # O filtro '$or' é comum em APIs NoSQL/vetoriais
                    filter_fb = {"$or": [{"tipo": "feedback"}, {"tipo": "feedback_attachment"}]}
                    # Busca por similaridade com filtro, limitando a k (ou um valor menor para feedback)
                    docs_potenciais_feedback = db_vectorstore.similarity_search(
                        query,
                        k=min(k, 10), # Busca menos documentos de feedback, foca nos mais relevantes
                        filter=filter_fb
                    )
                except Exception as e_filter:
                    st.warning(f"Não foi possível usar filtro de feedback ({e_filter}). Buscando feedback sem filtro específico...")
                    # Fallback sem filtro (pode trazer documentos não-feedback, serão filtrados depois)
                    docs_potenciais_feedback = db_vectorstore.similarity_search(query, k=min(k, 10))

                # Filtra e deduplica os resultados da busca de feedback
                seen_fb_content_hashes = set()
                for doc in docs_potenciais_feedback:
                    # Verifica se é realmente um documento de feedback (caso o filtro tenha falhado ou não usado)
                    if isinstance(doc.metadata, dict):
                        tipo = doc.metadata.get('tipo', '')
                        if tipo == 'feedback' or tipo == 'feedback_attachment':
                            content_hash = hash(doc.page_content) # Usa hash para deduplicação
                            if content_hash not in seen_fb_content_hashes:
                                # Marca como vindo da busca de feedback
                                doc.metadata['_retriever_name'] = "Feedback Search"
                                docs_feedback.append(doc)
                                seen_fb_content_hashes.add(content_hash)

                st.info(f"{len(docs_feedback)} documentos de feedback únicos encontrados.")
            else:
                st.warning("Retriever principal não possui 'vectorstore' ou é None. Busca de feedback pulada.")

        except Exception as e:
            st.warning(f"Erro durante a busca específica por feedback: {e}")


    # 3. Combinar e Deduplicar Resultados
    st.info("Combinando e deduplicando resultados...")
    unique_contents = set()
    combined_docs_final = []

    # Prioriza documentos de feedback primeiro (se houver)
    # Limita a proporção de feedback (ex: max 30% de k, ou um número fixo)
    max_feedbacks = max(1, int(k * 0.3)) if incluir_feedback else 0
    feedbacks_adicionados = 0

    for doc in docs_feedback:
         if feedbacks_adicionados < max_feedbacks:
              content_hash = hash(doc.page_content)
              if content_hash not in unique_contents:
                   combined_docs_final.append(doc)
                   unique_contents.add(content_hash)
                   feedbacks_adicionados += 1
         else:
              break # Atingiu o limite de feedback

    # Adiciona documentos dos outros retrievers, evitando duplicatas
    # Ordena pela relevância (assumindo que a busca já retorna ordenado) e pelo índice do retriever
    # Score pode não estar presente, usar fallback
    docs_normais_sorted = sorted(all_retrieved_docs, key=lambda d: (d.metadata.get('score', 1.0) if isinstance(d.metadata.get('score'), (int, float)) else 1.0, d.metadata.get('_retriever_index', 99)))

    for doc in docs_normais_sorted:
         content_hash = hash(doc.page_content)
         if content_hash not in unique_contents:
              # Adiciona até atingir o limite k
              if len(combined_docs_final) < k:
                   combined_docs_final.append(doc)
                   unique_contents.add(content_hash)
              else:
                   break # Atingiu o limite k

    st.info(f"Documentos RAG combinados e filtrados: {len(combined_docs_final)} (limite k={k}, {feedbacks_adicionados} de feedback)")

    # Opcional: Reordenar a lista final se necessário (ex: por relevância, se disponível e confiável)
    # combined_docs_final.sort(key=lambda d: d.metadata.get('score', 1.0), reverse=True) # Requer score confiável

    return combined_docs_final


# Histórico e feedback
def salvar_historico(pergunta, resposta, fontes_usadas_meta=None):
    """Salva a interação (pergunta, resposta, fontes) em um arquivo de texto."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        # Garante que o diretório base exista
        os.makedirs(os.path.dirname(HISTORICO_PATH), exist_ok=True)
        with open(HISTORICO_PATH, "a", encoding="utf-8") as f:
            f.write(f"--- Interaction @ {timestamp} ---\n")
            f.write(f"Pergunta:\n{pergunta}\n\n")
            f.write(f"Resposta Gerada:\n{resposta}\n\n")

            # Formata e escreve as fontes usadas, se disponíveis
            if fontes_usadas_meta and isinstance(fontes_usadas_meta, list):
                sources_set = set()
                fontes_formatadas = []
                for meta in fontes_usadas_meta:
                    if isinstance(meta, dict):
                        source_path = meta.get('source')
                        tipo = meta.get('tipo', 'Desconhecido')
                        # Adiciona outros metadados úteis, se existirem
                        chunk_idx = meta.get('start_index', meta.get('_chunk_id', 'N/A')) # Tenta obter índice do chunk
                        tipo_ext = meta.get('tipo_documento_ext', '')
                        id_ext = meta.get('id_documento_ext', '')

                        if source_path:
                            nome_base = os.path.basename(source_path)
                            # Constrói nome descritivo da fonte
                            display_parts = [f"Base: {nome_base}"]
                            if tipo_ext and tipo_ext not in ['DESCONHECIDO', 'ERRO_PARSE']: display_parts.append(f"TipoExt: {tipo_ext}")
                            if id_ext and id_ext != 'SEM_ID': display_parts.append(f"IDExt: {id_ext}")
                            if tipo != 'Desconhecido': display_parts.append(f"TipoMeta: {tipo}")
                            display_parts.append(f"ChunkIdx: {chunk_idx}")
                            display_name = " | ".join(display_parts)

                            # Evita duplicar o mesmo arquivo na lista de fontes, mas mostra todos os chunks
                            fontes_formatadas.append(f"- {display_name} | Path: {source_path}")
                            sources_set.add(nome_base) # Adiciona nome base ao set para contagem
                        else:
                             fontes_formatadas.append(f"- Fonte sem 'source' no metadado: {meta}")
                    else:
                        fontes_formatadas.append(f"- Metadado de fonte inválido (não é dict): {str(meta)}")

                if fontes_formatadas:
                    f.write(f"Fontes RAG Consultadas ({len(fontes_usadas_meta)} chunks de {len(sources_set)} arquivo(s) único(s)):\n")
                    f.write("\n".join(fontes_formatadas) + "\n")
                else:
                    f.write("Fontes RAG: Nenhuma metadado válido encontrado.\n")
            else:
                f.write("Fontes RAG: Nenhuma fornecida ou lista inválida.\n")

            f.write("-" * 40 + "\n\n") # Separador mais longo
        return timestamp # Retorna o timestamp para uso no feedback
    except Exception as e:
        st.error(f"Erro crítico ao salvar histórico em '{HISTORICO_PATH}': {e}")
        return None # Retorna None em caso de erro

def carregar_historico():
    """Carrega o conteúdo do arquivo de histórico."""
    if not os.path.exists(HISTORICO_PATH):
        return "Arquivo de histórico ainda não existe."
    if os.path.getsize(HISTORICO_PATH) == 0:
         return "Arquivo de histórico está vazio."
    try:
        with open(HISTORICO_PATH, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        st.error(f"Erro ao carregar histórico de '{HISTORICO_PATH}': {e}")
        return f"Erro ao ler arquivo de histórico: {e}"

def adicionar_comentario_rag(comentario, pergunta, resposta, classificacao, timestamp):
    """Salva um comentário/análise como um arquivo TXT e o indexa no RAG."""
    # Validação básica
    if not comentario or not isinstance(comentario, str) or not comentario.strip():
        st.info("Comentário vazio, nada a adicionar ao RAG.")
        return False

    # Cria um nome de arquivo seguro para o feedback
    safe_timestamp = timestamp.replace(':', '-').replace(' ', '_') if timestamp else datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    feedback_filename = f"analise_complementar_{safe_timestamp}.txt"
    feedback_source_path = os.path.join(FEEDBACK_CONTENT_DIR, feedback_filename)

    try:
        # Prepara metadados para este feedback
        metadata = {
            "source": feedback_source_path,
            "tipo": "feedback", # Marca como tipo feedback
            "pergunta_original": pergunta[:300] + ('...' if len(pergunta) > 300 else ''), # Limita tamanho
            "resposta_avaliada_trecho": resposta[:300] + ('...' if len(resposta) > 300 else ''), # Limita tamanho
            "classificacao": classificacao,
            "timestamp_feedback": timestamp or datetime.datetime.now().isoformat(),
            "comentario_curto": comentario[:150] + ('...' if len(comentario) > 150 else '') # Trecho do comentário
        }
        # Monta o conteúdo do arquivo TXT
        conteudo = f"""[ANÁLISE COMPLEMENTAR DO USUÁRIO - {metadata['timestamp_feedback']}]
Classificação da Resposta Original: {classificacao}
Pergunta Original (trecho): {metadata['pergunta_original']}
Resposta Avaliada (trecho): {metadata['resposta_avaliada_trecho']}

Comentário/Análise Detalhada:
{comentario}

[FIM DA ANÁLISE COMPLEMENTAR]"""

        # Salva o conteúdo no arquivo TXT
        try:
            os.makedirs(FEEDBACK_CONTENT_DIR, exist_ok=True) # Garante que o diretório exista
            with open(feedback_source_path, "w", encoding="utf-8") as f_fb:
                f_fb.write(conteudo)
            st.info(f"Análise complementar salva em: {feedback_filename}")
        except Exception as e_write:
            st.error(f"Não foi possível salvar o arquivo da análise complementar '{feedback_filename}': {e_write}")
            return False # Falha ao salvar o arquivo

        # Prepara para indexação: Cria um dict no formato esperado por dividir_documentos
        doc_para_dividir = [{"page_content": conteudo, "metadata": metadata}]

        # Divide o conteúdo do comentário (pode ser longo) usando a mesma config
        # Usar chunk_size menor para feedback pode ser útil se forem curtos
        chunks_comentario = dividir_documentos(doc_para_dividir, chunk_size=512, chunk_overlap=50)

        if not chunks_comentario:
            st.warning("Não foi possível gerar chunks para a análise complementar (pode ser muito curta). Não será indexada.")
            return False

        # Indexa os chunks gerados
        st.info(f"Indexando análise complementar ({len(chunks_comentario)} chunk(s)) no RAG...")
        chunks_indexed = indexar_documentos(chunks_comentario) # Passa os chunks

        if chunks_indexed > 0:
            st.success(f"✅ Análise complementar indexada com sucesso ({chunks_indexed} chunk(s)).")
            return True
        else:
            st.warning("⚠️ Não foi possível indexar a análise complementar no RAG após a divisão.")
            return False

    except Exception as e:
        st.error(f"Erro crítico ao processar e indexar análise complementar: {str(e)}")
        st.code(traceback.format_exc())
        return False

def salvar_feedback(pergunta, resposta, feedback_texto, classificacao, arquivos_feedback, timestamp):
    """Salva o feedback (CSV, comentário no RAG, anexos no RAG)."""
    feedback_timestamp = timestamp or datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    # Tenta converter para ISO format se não estiver, para consistência
    try:
        iso_timestamp = datetime.datetime.fromisoformat(feedback_timestamp.replace(' ', 'T')).isoformat()
    except ValueError:
         iso_timestamp = datetime.datetime.now().isoformat() # Fallback


    # 1. Salvar registro básico no CSV
    feedback_csv_data = {
        "timestamp": feedback_timestamp, # Mantém formato original no CSV por legibilidade
        "pergunta": pergunta,
        "resposta_trecho": resposta[:500] + ('...' if len(resposta) > 500 else ''), # Salva trecho da resposta
        "classificacao": classificacao,
        "comentario": feedback_texto if feedback_texto and isinstance(feedback_texto, str) else ""
        # Poderia adicionar contagem de anexos aqui se útil
        # "num_anexos": len(arquivos_feedback) if arquivos_feedback else 0
    }
    try:
        # Garante que o diretório base exista
        os.makedirs(os.path.dirname(FEEDBACK_PATH), exist_ok=True)
        # Verifica se o arquivo existe e está vazio para adicionar cabeçalho
        file_exists = os.path.exists(FEEDBACK_PATH)
        header = not file_exists or os.path.getsize(FEEDBACK_PATH) == 0
        # Usa 'utf-8-sig' para melhor compatibilidade com Excel em alguns sistemas
        pd.DataFrame([feedback_csv_data]).to_csv(FEEDBACK_PATH, mode='a', header=header, index=False, encoding='utf-8-sig')
        # st.info(f"Registro de feedback salvo em {FEEDBACK_PATH}")
    except Exception as e_csv:
        st.error(f"Erro ao salvar registro de feedback no CSV '{FEEDBACK_PATH}': {e_csv}")

    # 2. Indexar comentário/análise complementar no RAG (se houver)
    if feedback_texto and isinstance(feedback_texto, str) and feedback_texto.strip():
        # Passa o timestamp original para a função
        adicionar_comentario_rag(feedback_texto, pergunta, resposta, classificacao, feedback_timestamp)

    # 3. Processar e indexar anexos de feedback (se houver)
    if arquivos_feedback:
        st.info(f"Processando {len(arquivos_feedback)} anexo(s) de feedback para indexação...")
        docs_para_indexar_fb_total = []
        arquivos_fb_processados_ok = 0
        arquivos_fb_com_erro = 0

        for file_upload in arquivos_feedback:
            # Validação do objeto UploadedFile
            if not hasattr(file_upload, 'name') or not hasattr(file_upload, 'getvalue'):
                st.warning(f"Objeto de anexo inválido encontrado. Pulando.")
                arquivos_fb_com_erro += 1
                continue

            try:
                # Cria nome e caminho seguros para salvar o anexo
                safe_timestamp_fb = iso_timestamp.replace(':', '-').replace('.', '_') # Timestamp seguro para nome de arquivo
                # Limita tamanho do nome original e remove caracteres problemáticos
                safe_original_name = re.sub(r'[^\w\.-]', '_', file_upload.name[:50])
                nome_arquivo_fb = f"feedback_{safe_timestamp_fb}_{safe_original_name}"
                caminho_arquivo_fb = os.path.join(FEEDBACK_FILES_DIR, nome_arquivo_fb)

                # Salva o arquivo anexo fisicamente
                try:
                    os.makedirs(FEEDBACK_FILES_DIR, exist_ok=True) # Garante diretório
                    with open(caminho_arquivo_fb, "wb") as f_out:
                        f_out.write(file_upload.getvalue())
                    # st.info(f"Anexo de feedback '{file_upload.name}' salvo em {caminho_arquivo_fb}")
                except Exception as e_save_fb:
                     st.error(f"Erro ao salvar anexo de feedback '{file_upload.name}' em disco: {e_save_fb}")
                     arquivos_fb_com_erro += 1
                     continue # Pula para o próximo anexo se não puder salvar

                # Carrega o conteúdo do anexo salvo (TXT, PDF, DOCX)
                st.warning(f"DEBUG (salvar_feedback): Chamando carregar_arquivo para path='{caminho_arquivo_fb}', Nome original='{file_upload.name}'") # <-- Adicionar log
                docs_carregados_info = carregar_arquivo(caminho_arquivo_fb, is_attachment=True)

                # Verifica se o carregamento foi válido e tem conteúdo
                conteudo_valido_anexo = False
                if (docs_carregados_info and isinstance(docs_carregados_info, list) and
                        len(docs_carregados_info) > 0 and isinstance(docs_carregados_info[0], dict) and
                        docs_carregados_info[0].get("page_content") and
                        not docs_carregados_info[0]["page_content"].strip().startswith(("[Erro", "[Aviso", "[Conteúdo vazio"))):
                    conteudo_valido_anexo = True

                if conteudo_valido_anexo:
                    # Adiciona metadados específicos de feedback a cada documento carregado
                    for doc_info in docs_carregados_info:
                         if not isinstance(doc_info.get('metadata'), dict): doc_info['metadata'] = {}
                         doc_info['metadata'].update({
                             'source': caminho_arquivo_fb, # Caminho absoluto salvo
                             'tipo': 'feedback_attachment', # Marca como anexo de feedback
                             'origin': 'feedback_attachment',
                             'feedback_timestamp': iso_timestamp,
                             'original_filename': file_upload.name, # Nome original para referência
                             'classificacao_associada': classificacao # Classificação associada ao feedback
                         })
                         # Divide o conteúdo do anexo ANTES de adicionar à lista total
                         # Usar chunk_size menor para anexos pode ser razoável
                         chunks_anexo = dividir_documentos([doc_info], chunk_size=512, chunk_overlap=50)
                         if chunks_anexo:
                             docs_para_indexar_fb_total.extend(chunks_anexo)

                    arquivos_fb_processados_ok += 1
                else:
                    st.warning(f"Conteúdo inválido, vazio ou erro ao carregar anexo de feedback: {file_upload.name}. Não será indexado.")
                    arquivos_fb_com_erro += 1

            except Exception as e_file_fb:
                st.error(f"Erro inesperado ao processar anexo de feedback '{file_upload.name}': {str(e_file_fb)}")
                st.code(traceback.format_exc(limit=1))
                arquivos_fb_com_erro += 1

        # Indexa todos os chunks de anexos acumulados
        if docs_para_indexar_fb_total:
            st.info(f"Indexando {len(docs_para_indexar_fb_total)} chunk(s) de {arquivos_fb_processados_ok} anexo(s) de feedback válidos...")
            chunks_fb_indexados_total = indexar_documentos(docs_para_indexar_fb_total) # Indexa os chunks acumulados
            if chunks_fb_indexados_total > 0:
                st.success(f"✅ Conteúdo de {arquivos_fb_processados_ok} anexo(s) de feedback indexado ({chunks_fb_indexados_total} chunks).")
            else:
                st.warning("⚠️ Nenhum chunk dos anexos de feedback pôde ser indexado.")
        elif arquivos_feedback: # Se havia arquivos, mas nenhum gerou chunks válidos
             st.info("Nenhum anexo de feedback continha conteúdo válido ou passível de divisão para indexar.")

        if arquivos_fb_com_erro > 0:
            st.warning(f"{arquivos_fb_com_erro} anexo(s) de feedback encontraram erros durante o processamento.")


# Função para verificar se o DB está vazio
def verificar_db_vazio():
    """Verifica se o diretório do ChromaDB existe e se a coleção contém itens."""
    persist_directory = PERSIST_DIR
    # Se o diretório nem existe, está vazio
    if not os.path.exists(persist_directory) or not os.listdir(persist_directory):
        return True

    global embeddings
    if embeddings is None:
        st.warning("Embeddings não inicializados, não é possível verificar o conteúdo do DB.")
        return True # Assume vazio se não pode verificar

    db_check = None
    try:
        # Tenta carregar a instância do DB
        db_check = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
        # Tenta contar itens. Isso pode falhar se a coleção não existir.
        count = db_check._collection.count()
        return count == 0
    except Exception as e:
        # Se houver erro (ex: coleção não existe, DB corrompido), trata como vazio/problemático
        st.warning(f"Erro ao verificar contagem de itens no DB: {e}. Tratando como vazio/problemático.")
        return True
    finally:
         # Garante que qualquer conexão seja fechada (se aplicável ao ChromaDB local)
         # A biblioteca Chroma geralmente gerencia isso, mas é uma boa prática pensar sobre.
         if hasattr(db_check, '_client') and hasattr(db_check._client, 'stop'):
              # db_check._client.stop() # Descomentar se souber que isso é necessário/seguro
              pass

# Função de inicialização com exemplos
def inicializar_chromadb_exemplo(db_instance=None):
    """Adiciona alguns documentos de exemplo ao ChromaDB se estiver vazio."""
    global embeddings
    if embeddings is None:
        st.error("Embeddings não inicializados. Não é possível adicionar exemplos.")
        return False

    db_exemplo = None
    try:
        # Obtém a instância do DB (ou carrega se não fornecida)
        db_exemplo = db_instance if db_instance else carregar_chroma()
        if db_exemplo is None:
            st.error("Falha ao obter instância do DB para adicionar exemplos.")
            return False

        # Verifica se já contém dados (exceto o dummy inicial ou exemplos anteriores)
        try:
            count = db_exemplo._collection.count()
            results_existentes = db_exemplo.get(where={"$or": [{"tipo": "texto_exemplo"}, {"source":"dummy"}]})
            num_existentes = len(results_existentes.get('ids', []))
            # Se a contagem total for maior que o número de exemplos/dummies, assume que já tem dados reais
            if count > num_existentes:
                 st.info(f"Base de dados já contém {count - num_existentes} item(ns) além dos exemplos/dummies. Não adicionando exemplos novamente.")
                 return True
        except Exception as e_check:
             # Se a contagem ou get falhar, assume que pode estar vazio ou corrompido, tenta adicionar
             st.warning(f"Não foi possível verificar conteúdo existente antes de adicionar exemplos ({e_check}). Tentando adicionar...")


        st.info("Base de dados está vazia ou contém apenas exemplos. Adicionando/Verificando exemplos...")

        # Textos de exemplo
        textos_exemplo = [
            ("Perícia Android: Ferramentas como UFED e XRY são comumente usadas para extração lógica, de sistema de arquivos e física. A extração física oferece a imagem mais completa, mas nem sempre é possível.", {"tipo_doc": "Nota Técnica", "topico": "Perícia Móvel"}),
            ("ISO 27037: Esta norma estabelece diretrizes para identificação, coleta, aquisição e preservação de evidência digital potencial. Garante a admissibilidade e integridade.", {"tipo_doc": "Norma ISO", "topico": "Boas Práticas"}),
            ("Cadeia de Custódia (CPP): Os artigos 158-A a 158-F do Código de Processo Penal brasileiro detalham os procedimentos para manter a integridade e rastreabilidade da prova digital, desde a coleta até o descarte.", {"tipo_doc": "Legislação", "topico": "Cadeia de Custódia"}),
            ("Análise Forense de Logs: Logs de sistema (syslog, EventLog), logs de firewall e logs de aplicação são fontes cruciais para reconstruir eventos e identificar atividades suspeitas.", {"tipo_doc": "Procedimento", "topico": "Análise Forense"}),
            ("Hashing e Integridade: Algoritmos de hash como SHA-256 e MD5 geram uma assinatura digital única para arquivos ou imagens forenses. Comparar hashes antes e depois da cópia verifica se houve alteração.", {"tipo_doc": "Conceito", "topico": "Integridade de Dados"})
        ]

        docs_exemplo_info = []
        # Garante que o diretório de exemplos exista
        os.makedirs(EXEMPLOS_INTERNOS_DIR, exist_ok=True)

        for i, (texto, meta_extra) in enumerate(textos_exemplo):
            # Cria um caminho "fictício" dentro da pasta de exemplos para o source
            source_path = os.path.join(EXEMPLOS_INTERNOS_DIR, f"exemplo_{i+1}.txt")
            metadata = {"source": source_path, "tipo": "texto_exemplo"} # Marca o tipo
            metadata.update(meta_extra) # Adiciona metadados específicos do exemplo
            metadata["indexed_at"] = datetime.datetime.now().isoformat()

            # Salva o exemplo em um arquivo (opcional, mas bom para consistência)
            try:
                with open(source_path, 'w', encoding='utf-8') as f_ex:
                    f_ex.write(f"# Exemplo {i+1} - {meta_extra.get('topico', '')}\n\n{texto}")
            except Exception as e_write_ex:
                st.warning(f"Erro ao salvar arquivo de exemplo {source_path}: {e_write_ex}")

            # Pré-processa o texto do exemplo
            texto_processado = preprocessar_texto(texto)
            if texto_processado:
                # Cria o dict no formato esperado por dividir_documentos
                docs_exemplo_info.append({"page_content": texto_processado, "metadata": metadata})

        if not docs_exemplo_info:
            st.error("Nenhum exemplo válido foi processado. Algo deu errado.")
            return False

        # Divide os exemplos em chunks (importante para RAG funcionar bem)
        # Usar chunk_size menor pode fazer sentido para exemplos curtos
        chunks_exemplos = dividir_documentos(docs_exemplo_info, chunk_size=256, chunk_overlap=30)
        if not chunks_exemplos:
            st.error("Nenhum chunk foi gerado para os exemplos. Conteúdo pode ser muito curto.")
            return False

        # Indexa os chunks dos exemplos
        st.info(f"Indexando {len(chunks_exemplos)} chunks de exemplo...")
        chunks_added = indexar_documentos(chunks_exemplos) # Indexa os chunks

        if chunks_added > 0:
            st.success(f"Exemplos adicionados/atualizados com sucesso ({chunks_added} chunks).")
            return True
        else:
            st.error("Falha ao indexar os chunks dos exemplos.")
            return False

    except Exception as e:
        st.error(f"Erro crítico durante a inicialização com exemplos: {str(e)}")
        st.code(traceback.format_exc())
        return False
    finally:
         # Garante que qualquer conexão seja fechada (se aplicável)
         if hasattr(db_exemplo, '_client') and hasattr(db_exemplo._client, 'stop'):
             # db_exemplo._client.stop() # Descomentar se souber que é necessário/seguro
             pass

# --- Interface Streamlit ---
def main():
    st.set_page_config(page_title="Professor de Direito Digital", layout="wide", initial_sidebar_state="expanded")
    st.title("Professor de Direito Digital 🔎🤖")
    st.caption("Assistente IA para ensino de Direito Digital.")

    # Verifica inicialização essencial (OpenAI embeddings e cliente)
    global embeddings, openai_client_instance
    if embeddings is None or openai_client_instance is None:
         st.error("Inicialização da API OpenAI falhou. Verifique a chave API e a conexão com a internet. A aplicação não pode continuar.")
         st.stop()

    # Verifica e inicializa o DB na primeira execução da sessão
    if 'system_initialized' not in st.session_state:
        st.session_state['system_initialized'] = False
        with st.spinner("Verificando base de dados RAG..."):
            db_vazio = verificar_db_vazio()

        if db_vazio:
            st.warning("⚠️ A base de dados RAG parece estar vazia ou é a primeira execução.")
            if st.button("➕ Inicializar Base com Exemplos"):
                with st.spinner("Adicionando exemplos à base de dados..."):
                    if inicializar_chromadb_exemplo():
                        st.success("✅ Sistema inicializado com exemplos! Base pronta para uso.")
                        st.session_state['system_initialized'] = True
                        st.rerun() # Recarrega para refletir o estado
                    else:
                        st.error("❌ Falha ao inicializar com exemplos. Verifique os logs e permissões.")
                        # Não marca como inicializado se falhar
            else:
                st.info("A base está vazia. Você pode inicializar com exemplos acima ou usar 'Indexar Textos' na barra lateral para adicionar seus próprios documentos.")
                # Decide se permite continuar com base vazia
                # st.session_state['system_initialized'] = True # Descomentar para permitir uso mesmo vazio
        else:
            st.success("✅ Base de dados RAG encontrada e contém dados.")
            st.session_state['system_initialized'] = True

    # Se o sistema não foi inicializado (DB vazio e usuário não clicou no botão),
    # pode parar ou mostrar mensagem mais proeminente.
    if not st.session_state.get('system_initialized', False):
         st.warning("Sistema não inicializado ou base vazia. Use as opções na barra lateral para preparar a base de dados.")
         # st.stop() # Descomentar se quiser impedir o uso sem inicialização

    # Inicializa/mantém estado da sessão para a conversa atual
    default_session_state = {
        'query': "",
        'resposta': "",
        'fontes_usadas': [], # Armazena metadados das fontes da última resposta
        'timestamp': None, # Timestamp da última resposta
        'current_query_for_feedback': "", # Query associada ao feedback atual
        'feedback_submitted': False, # Flag para controlar exibição do form de feedback
        'cached_fontes': None, # Cache da lista de fontes indexadas
        'df_feedback_audit': None, # DataFrame para auditoria de feedback
        'feedback_loaded': False # Flag para auditoria de feedback
        }
    for key, default_value in default_session_state.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

    # --- Sidebar ---
    with st.sidebar:
        st.header("Configurações e Gerenciamento")
        st.markdown("---") # Separador visual

        # Seleção de Modelo
        st.subheader("⚙️ Modelo de IA")
        # Lista de modelos disponíveis (pode ser dinâmica no futuro)
        model_options_map = {
            "gpt-4o": "GPT-4o (Recomendado)",
            "gpt-4-turbo": "GPT-4 Turbo",
            # "gpt-4": "GPT-4",
            # "gpt-3.5-turbo": "GPT-3.5 Turbo",
            # Adicionar outros modelos se disponíveis e testados
            "gpt-4.1": "GPT-4.1 Mini",
            # "text-davinci-003": "Davinci (Legado)"
        }
        available_models = list(model_options_map.keys())
        # Preferência de modelo padrão
        default_model_preference = ["gpt-4o", "gpt-4-turbo", "gpt-4", "gpt-3.5-turbo"]
        # Encontra o melhor modelo padrão disponível na lista
        default_model_key = next((m for m in default_model_preference if m in available_models), available_models[0])
        # Widget de seleção
        model_choice = st.selectbox(
            "Modelo:",
            options=available_models,
            format_func=lambda x: model_options_map.get(x, x), # Mostra nome amigável
            index=available_models.index(default_model_key),
            key="model_choice_selector", # Chave para o estado da sessão
            help="Escolha o modelo de linguagem para gerar as respostas."
        )
        # Mostra limites do modelo selecionado
        model_max_context = obter_limite_contexto(model_choice)
        model_max_output = obter_limite_output(model_choice)
        st.caption(f"Contexto Total: {model_max_context:,} tokens | Limite Saída: {model_max_output:,} tokens".replace(',', '.'))
        st.markdown("---")

        # Parâmetros da Resposta
        st.subheader("📝 Parâmetros da Resposta")
        model_temperature = st.slider(
            "Criatividade (Temperatura):",
            min_value=0.0, max_value=1.5, value=0.2, step=0.05,
            key="temperature_slider",
            help="Valores mais baixos geram respostas mais focadas e determinísticas. Valores mais altos geram respostas mais criativas e variadas."
        )
        # Define o máximo de tokens desejados, limitado pelo output do modelo
        # Ajustado valor padrão para ser um pouco menor, considerando multi-partes
        default_response_tokens = min(1500, model_max_output)
        desired_max_tokens = st.number_input(
            "Máximo de Tokens por Resposta (Parte):",
            min_value=100, max_value=model_max_output, value=default_response_tokens, step=100,
            key="desired_max_tokens", # Chave para o estado
            help=f"Define o limite máximo de tokens para *cada parte* da resposta gerada. O limite real do modelo é {model_max_output}."
        )
        use_feedback_in_context = st.checkbox(
            "Considerar feedback/análises anteriores nas buscas RAG",
            value=True,
            key="use_feedback_checkbox",
            help="Se marcado, o sistema tentará buscar também por feedbacks e análises complementares relevantes."
        )
        st.markdown("---")

        # Gerenciamento da Base RAG
        st.subheader("📚 Gerenciamento da Base RAG")
        with st.expander("Indexação e Diagnóstico", expanded=False):
            st.markdown("**Indexação Manual da Pasta OCR**")
            recursive_index = st.checkbox("Incluir subpastas ao indexar", value=True, key="recursive_index_chk")
            if st.button("⚙️ Indexar Textos Processados (OCR)", key="index_ocr_btn", help=f"Busca e indexa arquivos .txt da pasta configurada para OCR ({os.path.basename(OCR_OUTPUT_DIR)})."):
                if not os.path.isdir(OCR_OUTPUT_DIR):
                    st.error(f"Erro: Diretório de textos OCR '{OCR_OUTPUT_DIR}' não encontrado. Verifique a configuração.")
                else:
                    with st.spinner(f"Indexando .txt em '{os.path.basename(OCR_OUTPUT_DIR)}'... Isso pode levar alguns minutos."):
                        proc_ok, idx_count = indexar_pasta_completa(OCR_OUTPUT_DIR, recursive_index)
                    # Feedback após indexação
                    if idx_count > 0:
                        st.success(f"✅ Indexação concluída! {proc_ok} arquivos válidos processados, {idx_count} chunks adicionados/atualizados.")
                        st.session_state['cached_fontes'] = None # Limpa cache da lista de fontes
                        st.rerun() # Atualiza a lista na UI
                    elif proc_ok > 0:
                        st.warning(f"Indexação concluída, mas nenhum chunk novo foi adicionado ({proc_ok} arquivos processados).")
                    else:
                        st.error("❌ Falha na indexação. Nenhum arquivo válido processado ou nenhum chunk gerado.")

            st.markdown("**Fontes Indexadas (Amostra)**")
            # Atualiza cache se necessário
            if st.session_state.get('cached_fontes') is None:
                 with st.spinner("Atualizando lista de fontes indexadas..."):
                      st.session_state['cached_fontes'] = listar_arquivos_indexados()

            # Exibe a lista em cache
            if st.session_state['cached_fontes']:
                st.text_area("Fontes Encontradas (Nomes Base):",
                             "\n".join(st.session_state['cached_fontes']),
                             height=150, key="fontes_list_area", disabled=True,
                             help="Lista de nomes de arquivos encontrados nos metadados da base RAG (exclui exemplos e feedback).")
            else:
                st.info("Nenhuma fonte principal encontrada na base de dados (ou a base está vazia).")

            if st.button("🔄 Atualizar Lista de Fontes", key="refresh_sources_btn"):
                st.session_state['cached_fontes'] = None # Força recarregamento na próxima renderização
                st.rerun()

            st.markdown("**Diagnóstico e Reparo**")
            if st.button("🩺 Diagnosticar ChromaDB", key="diag_chroma_btn"):
                diagnosticar_chromadb()

            st.error("🚨 Ação Perigosa Abaixo!") # Destaque para o botão de reparo
            if st.button("⚠️ Reparar ChromaDB (Recriar Base)", key="repair_chroma_btn", help="APAGA COMPLETAMENTE a base de dados atual e cria uma nova. Use com cuidado!"):
                 st.warning("🚨 ATENÇÃO: Esta ação é irreversível e apagará TODOS os dados indexados (documentos, exemplos, feedback).")
                 # Confirmação dupla para segurança
                 with st.form("repair_confirm_form"):
                     st.markdown("**Digite 'REPARAR' para confirmar:**")
                     confirmation_text = st.text_input("Confirmação:", key="repair_confirm_text", type="password") # Usa tipo password para esconder
                     submitted = st.form_submit_button("Confirmar e Recriar Base")
                     if submitted:
                         if confirmation_text.strip().upper() == "REPARAR":
                             st.info("Iniciando processo de recriação...")
                             try:
                                 # 1. Backup (Opcional, mas recomendado)
                                 backup_dir = f"{PERSIST_DIR}_backup_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
                                 st.info(f"Tentando criar backup em: {backup_dir}")
                                 try:
                                     if os.path.exists(PERSIST_DIR):
                                         shutil.copytree(PERSIST_DIR, backup_dir)
                                         st.success(f"Backup da base antiga criado em: {backup_dir}")
                                     else:
                                          st.info("Diretório original não existe, backup pulado.")
                                 except Exception as backup_err:
                                     st.error(f"Falha ao criar backup: {backup_err}. O processo de recriação continuará, mas sem backup.")

                                 # 2. Remover diretório antigo
                                 st.info(f"Removendo diretório antigo: {PERSIST_DIR}")
                                 if os.path.exists(PERSIST_DIR):
                                     shutil.rmtree(PERSIST_DIR)

                                 # 3. Recriar diretório e inicializar DB
                                 st.info(f"Recriando diretório e inicializando nova base...")
                                 os.makedirs(PERSIST_DIR, exist_ok=True)
                                 db = carregar_chroma() # Tenta criar nova instância
                                 if db:
                                     st.success("✅ ChromaDB recriado com sucesso!")
                                     # Opcional: Adicionar exemplos à nova base
                                     st.info("Adicionando exemplos à nova base...")
                                     inicializar_chromadb_exemplo(db)
                                     st.session_state['cached_fontes'] = None # Limpa cache
                                     st.session_state['system_initialized'] = True # Marca como inicializado
                                     st.rerun() # Recarrega a aplicação
                                 else:
                                     st.error("❌ Falha crítica ao inicializar o novo ChromaDB após a remoção.")

                             except Exception as repair_err:
                                 st.error(f"Erro durante o processo de reparo/recriação: {repair_err}")
                                 st.code(traceback.format_exc())
                         else:
                             st.error("Confirmação incorreta. A recriação foi cancelada.")
        st.markdown("---")

        # Auditoria
        st.subheader("📊 Auditoria")
        with st.expander("Visualizar Feedbacks Salvos (CSV)"):
            if st.button("Carregar Histórico de Feedbacks (CSV)", key="load_feedback_btn"):
                if os.path.exists(FEEDBACK_PATH) and os.path.getsize(FEEDBACK_PATH) > 0:
                    try:
                        # Tenta ler com utf-8-sig primeiro
                        df_feedback = pd.read_csv(FEEDBACK_PATH, encoding='utf-8-sig')
                        st.session_state['df_feedback_audit'] = df_feedback
                        st.session_state['feedback_loaded'] = True
                        st.success(f"{len(st.session_state['df_feedback_audit'])} registros de feedback carregados do CSV.")
                    except UnicodeDecodeError:
                         try: # Fallback para utf-8
                              st.info("Falha com utf-8-sig, tentando utf-8...")
                              df_feedback = pd.read_csv(FEEDBACK_PATH, encoding='utf-8')
                              st.session_state['df_feedback_audit'] = df_feedback
                              st.session_state['feedback_loaded'] = True
                              st.success(f"{len(st.session_state['df_feedback_audit'])} registros de feedback carregados do CSV (utf-8).")
                         except Exception as e_read:
                              st.error(f"Erro ao ler CSV de feedback com utf-8: {e_read}")
                              st.session_state['feedback_loaded'] = False
                    except Exception as e:
                        st.error(f"Erro inesperado ao ler CSV de feedback: {e}")
                        st.session_state['feedback_loaded'] = False
                else:
                    st.warning(f"Arquivo de feedback '{os.path.basename(FEEDBACK_PATH)}' não encontrado ou está vazio.")
                    st.session_state['feedback_loaded'] = False
                    if 'df_feedback_audit' in st.session_state: del st.session_state['df_feedback_audit'] # Limpa df antigo

            # Exibe o DataFrame se carregado
            if st.session_state.get('feedback_loaded', False) and 'df_feedback_audit' in st.session_state and isinstance(st.session_state['df_feedback_audit'], pd.DataFrame) and not st.session_state['df_feedback_audit'].empty:
                df_audit = st.session_state['df_feedback_audit']
                st.metric("Total de Feedbacks Registrados", len(df_audit))
                # Mostra contagem por classificação, se a coluna existir
                if 'classificacao' in df_audit.columns:
                    st.write("**Distribuição por Classificação:**")
                    st.bar_chart(df_audit['classificacao'].value_counts())
                else:
                     st.info("Coluna 'classificacao' não encontrada no CSV.")
                st.write("**Registros de Feedback:**")
                # Mostra o DataFrame com altura limitada e largura do container
                st.dataframe(df_audit, height=300, use_container_width=True)
            elif st.session_state.get('feedback_loaded', False): # Carregou mas está vazio
                 st.info("O arquivo de feedback foi carregado, mas está vazio.")

        st.markdown("---")

        # Teste de Conexão API
        st.subheader("📡 Teste de Conexão OpenAI")
        if st.button("Testar Conexão API", key="test_api_btn", help="Verifica se a chave API configurada permite a comunicação com o modelo selecionado."):
            modelo_sel_teste = st.session_state.get('model_choice_selector', default_model_key)
            with st.spinner(f"Testando conexão com {modelo_sel_teste}..."):
                ok, msg = testar_openai(modelo_sel_teste)
            if ok:
                st.success(f"✅ {msg}")
            else:
                st.error(f"❌ {msg}")
        st.markdown("---")


        # Informações Técnicas
        st.subheader("ℹ️ Informações Técnicas")
        with st.expander("Ver detalhes do ambiente"):
            st.text(f"Versão Python: {sys.version.split()[0]}")
            st.text(f"Diretório Base Configurado: {BASE_FONTES}")
            st.text(f"Diretório Textos OCR: {OCR_OUTPUT_DIR}")
            st.text(f"Diretório ChromaDB: {PERSIST_DIR}")
            st.text(f"Diretório Script: {script_dir}")
            st.markdown("**Versões das Bibliotecas Principais:**")
            # Define o dicionário DENTRO do expander para garantir que os imports ocorreram
            libs_to_check = {
                'streamlit': st,
                'langchain': langchain, # Agora deve funcionar com o import adicionado
                'chromadb': chromadb,
                'openai': openai_client_instance, # Mostra a instância do cliente (não tem __version__)
                'pandas': pd,
                'PyPDF2': PyPDF2,
                'python-docx': Document, # Tratamento especial abaixo
                'chardet': chardet,
                'tiktoken': tiktoken
            }
            versions_text = []
            for name, lib_module in libs_to_check.items():
                version = "N/A"
                try:
                    if name == 'python-docx':
                        # Importa o pacote raiz para obter a versão
                        import docx as docx_lib
                        if hasattr(docx_lib, '__version__'):
                            version = f"v{docx_lib.__version__}"
                        else: version = "(Versão não encontrada)"
                    elif name == 'openai':
                         # O cliente não tem __version__, mas podemos tentar pegar do pacote base
                         try:
                              import openai as openai_base_lib
                              if hasattr(openai_base_lib, '__version__'):
                                   version = f"v{openai_base_lib.__version__}"
                              else: version = "(Cliente carregado, versão pacote N/A)"
                         except ImportError: version = "(Pacote openai não importado?)"
                    elif hasattr(lib_module, '__version__'):
                        version = f"v{lib_module.__version__}"
                    else:
                         version = "(Versão não encontrada)"
                except NameError:
                     version = "(Não importado)"
                except Exception as e_ver:
                     version = f"(Erro: {e_ver})"
                versions_text.append(f"- {name}: {version}")

            st.text('\n'.join(versions_text))

    # --- Layout Principal ---
    col_main, col_uploader = st.columns([3, 1]) # Coluna principal mais larga

    with col_main:
        st.subheader("💬 Consulta e Análise")
        # Área de texto para a pergunta
        query = st.text_area(
            "Sua pergunta:",
            value=st.session_state.query, # Mantém o valor entre runs
            height=150,
            key="query_input", # Chave para o estado
            placeholder="Descreva o processo de coleta de dados em um dispositivo Android usando UFED, mencionando a importância da cadeia de custódia e as diretrizes da ISO 27037."
        )
        # Atualiza o estado da query a cada alteração (útil se precisar reagir a mudanças)
        # Isso pode causar reruns frequentes se não for necessário. Remover se causar lentidão.
        # st.session_state.query = query

        # Botões de ação
        b_col1, b_col2 = st.columns(2)
        analisar_btn = b_col1.button("🔍 Analisar com RAG", use_container_width=True, type="primary", key="analisar_btn")
        limpar_btn = b_col2.button("🗑️ Limpar Consulta", use_container_width=True, key="limpar_btn")

    with col_uploader:
        st.subheader("📎 Anexar (Temporário)")
        # Uploader de arquivos
        files_uploaded = st.file_uploader(
            "Anexar arquivos (PDF, DOCX, TXT) para esta consulta:",
            accept_multiple_files=True,
            key="file_uploader", # Chave para o estado
            help="Os arquivos anexados aqui são usados APENAS para a consulta atual e não são indexados permanentemente na base principal."
        )

    # --- Lógica Limpeza ---
    if limpar_btn:
        # Reseta os campos relevantes da sessão para o padrão
        st.session_state.query = ""
        st.session_state.resposta = ""
        st.session_state.fontes_usadas = []
        st.session_state.timestamp = None
        st.session_state.current_query_for_feedback = ""
        st.session_state.feedback_submitted = False
        # Limpar o file_uploader requer uma abordagem diferente ou apenas confiar no rerun
        st.success("Campos da consulta limpos.")
        # Força recarregar a UI para limpar o uploader visualmente
        # Cuidado: st.rerun() pode ter efeitos colaterais dependendo da lógica.
        st.rerun()

    # --- Lógica Principal Análise ---
    # Executa apenas se o botão de analisar for clicado
    if analisar_btn:
        # Atualiza a query no estado ANTES de iniciar a análise
        st.session_state.query = query
        if not st.session_state.query: # Usa a query do estado
            st.error("❗ Por favor, digite uma pergunta para análise.")
        # Verifica se o sistema está inicializado (DB pronto)
        elif not st.session_state.get('system_initialized', False):
             st.error("❗ O sistema ainda não está pronto. Por favor, inicialize a base de dados ou indexe documentos na barra lateral.")
        else:
            # Reseta estado da resposta anterior antes de iniciar nova análise
            st.session_state.update({'resposta': "", 'fontes_usadas': [], 'timestamp': None, 'feedback_submitted': False, 'current_query_for_feedback': st.session_state.query})

            # Placeholder para UI de progresso
            analysis_placeholder = st.empty()
            with analysis_placeholder.container():
                progress_bar = st.progress(0.0)
                status_text = st.info("🧠 Preparando para análise...")

            temp_paths_to_clean = [] # Lista para guardar caminhos de arquivos temporários a limpar
            docs_temp_chunks = [] # Lista para guardar chunks dos arquivos temporários

            try:
                # ===============================================================
                # == INÍCIO DO BLOCO DE ANÁLISE MULTI-RESPOSTA ==
                # ===============================================================

                # Etapa 1: Processar anexos temporários (Uploads)
                progress_bar.progress(0.05, text="Verificando anexos...")
                if files_uploaded:
                    status_text.info(f"📄 Processando {len(files_uploaded)} anexo(s) temporário(s)...")
                    os.makedirs(SUBPASTA_TEMP, exist_ok=True) # Garante que a pasta temp existe
                    for i, f in enumerate(files_uploaded):
                        # Atualiza progresso e status
                        progress_percentage = 0.05 + ((i + 1) / len(files_uploaded)) * 0.15 # Progresso de 5% a 20%
                        progress_bar.progress(progress_percentage, text=f"Lendo anexo: {f.name}")

                        # Salva anexo temporariamente
                        tmp_path = os.path.join(SUBPASTA_TEMP, f"{datetime.datetime.now().timestamp()}_{f.name}")
                        temp_paths_to_clean.append(tmp_path) # Adiciona à lista para limpeza posterior
                        try:
                            with open(tmp_path, "wb") as fp:
                                fp.write(f.getvalue())

                            # Carrega e pré-processa o conteúdo do anexo salvo
                            loaded_info = carregar_arquivo(tmp_path, is_attachment=True)

                            # Verifica se carregou com sucesso e tem conteúdo válido
                            conteudo_valido_anexo = False
                            if (loaded_info and isinstance(loaded_info, list) and len(loaded_info) > 0 and
                                    isinstance(loaded_info[0], dict) and loaded_info[0].get("page_content") and
                                    not loaded_info[0]["page_content"].strip().startswith(("[Erro", "[Aviso","[Conteúdo vazio"))):
                                conteudo_valido_anexo = True

                            if conteudo_valido_anexo:
                                # Adiciona metadados indicando origem temporária
                                for item in loaded_info:
                                    if not isinstance(item.get('metadata'), dict): item['metadata'] = {}
                                    item['metadata'].update({'origin': 'temporary_attachment', 'original_filename': f.name})

                                # Divide o conteúdo do anexo em chunks
                                # Usar chunk_size menor para anexos pode ser útil
                                chunks_anexo_temp = dividir_documentos(loaded_info, chunk_size=512, chunk_overlap=50)
                                if chunks_anexo_temp:
                                    docs_temp_chunks.extend(chunks_anexo_temp) # Acumula os chunks
                            else:
                                st.warning(f"⚠️ Anexo temporário '{f.name}' inválido, vazio ou com erro no carregamento. Será ignorado.")
                        except Exception as e_tmp:
                            st.warning(f"⚠️ Erro ao processar anexo temporário '{f.name}': {e_tmp}")
                    progress_bar.progress(0.20, text="Anexos temporários processados.")
                else:
                    progress_bar.progress(0.20, text="Nenhum anexo temporário fornecido.") # Pula para 20%

                # Etapa 2: Preparar Retrievers (Principal e Temporário)
                status_text.info("🔍 Preparando busca RAG (Base Principal)...")
                progress_bar.progress(0.25, text="Carregando base de dados principal...")
                db_principal = carregar_chroma() # Carrega o DB persistente
                if db_principal is None:
                    # Erro crítico, não pode continuar sem DB principal
                    status_text.error("❌ Falha crítica ao carregar a base de dados principal. Análise abortada.")
                    progress_bar.progress(1.0)
                    st.stop() # Interrompe a execução

                # Cria retriever para o DB principal
                # k_ind: número de chunks a buscar SÓ da base principal (ajuste conforme necessidade)
                k_ind = 100 # Exemplo: busca até 50 chunks da base principal
                retrievers = [db_principal.as_retriever(search_kwargs={"k": k_ind})]
                st.info(f"Retriever da base principal pronto (k={k_ind}).")

                # Cria retriever para os anexos temporários (se houver chunks deles)
                if docs_temp_chunks:
                    progress_bar.progress(0.30, text="Criando índice temporário para anexos...")
                    status_text.info("⚙️ Criando índice em memória para anexos...")
                    try:
                        # 1. Instancia um cliente ChromaDB efêmero (em memória, local)
                        from chromadb.config import Settings

                        # Inicializa EphemeralClient com configurações explícitas para garantir isolamento
                        temp_client = chromadb.EphemeralClient() # Usa os padrões do cliente efêmero

                        # 2. Cria o ChromaDB temporário usando o cliente efêmero explicitamente
                        # Em vez de from_documents, inicializa Chroma e depois adiciona os documentos
                        db_temp = Chroma(
                        client=temp_client, # Passa o cliente efêmero
                        embedding_function=embeddings,
                        # Pode definir um nome de coleção único se quiser, mas EphemeralClient geralmente isola bem
                        collection_name=f"temp_collection_{datetime.datetime.now().timestamp()}"
                        )
                        # Adiciona os documentos à coleção recém-criada
                        db_temp.add_documents(docs_temp_chunks)

                        # k_temp: número de chunks a buscar SÓ dos anexos (geralmente menor)
                        k_temp = 30 # Exemplo: busca até 30 chunks dos anexos
                        retrievers.append(db_temp.as_retriever(search_kwargs={"k": k_temp}))
                        st.info(f"Retriever temporário para anexos pronto (k={k_temp}). Total retrievers: {len(retrievers)}")
                        progress_bar.progress(0.40, text="Índice temporário criado.")
                    except Exception as e_mem_db:
                        st.error(f"Erro ao criar índice temporário em memória para anexos: {str(e_mem_db)}. Anexos podem não ser consultados.")
                        progress_bar.progress(0.40) # Continua sem o retriever temp se falhar
                else:
                    progress_bar.progress(0.40) # Pula para 40% se não houver anexos

                # Etapa 3: Buscar Documentos (RAG) usando todos os retrievers
                status_text.info("📚 Combinando resultados da busca RAG...")
                # k_rag: número MÁXIMO de chunks a serem combinados e enviados ao LLM
                k_rag = 130 # Ajuste este valor para balancear contexto e tamanho da resposta
                progress_bar.progress(0.50, text=f"Buscando e combinando até {k_rag} chunks RAG...")

                # Chama a função que busca em todos os retrievers e combina/filtra
                docs_combinados = combinar_resultados(
                    retrievers,
                    st.session_state.query, # Usa a query do estado
                    k=k_rag, # Limite final de chunks
                    incluir_feedback=st.session_state.use_feedback_checkbox # Usa a configuração da UI
                )

                # Preparar contexto formatado e metadados UMA VEZ
                contexto_formatado = "[Contexto RAG não encontrado ou vazio]"
                fontes_usadas_meta_completas = []
                if docs_combinados:
                    contexto_formatado = "" # Reseta para construir
                    seen_sources_content_hashes = set() # Para evitar duplicatas exatas de conteúdo no prompt
                    for i, doc in enumerate(docs_combinados):
                        # Usa hash do conteúdo para verificar duplicação (mesmo se metadados diferentes)
                        content_hash = hash(doc.page_content)
                        if content_hash not in seen_sources_content_hashes:
                            # Adiciona metadados à lista para salvar no histórico
                            if isinstance(doc.metadata, dict):
                                fontes_usadas_meta_completas.append(doc.metadata)
                            else: # Caso inesperado de metadado inválido
                                fontes_usadas_meta_completas.append({"source": "Metadado Inválido", "chunk_idx": i})

                            seen_sources_content_hashes.add(content_hash)

                            # Formata a exibição da fonte para o prompt do LLM
                            try:
                                fonte_display = f"Fonte RAG #{i+1}" # Fallback
                                meta = doc.metadata if isinstance(doc.metadata, dict) else {}
                                # Constrói um nome mais descritivo para a fonte
                                src = meta.get('source')
                                tipo = meta.get('tipo', '')
                                orig_fname = meta.get('original_filename')
                                tipo_ext = meta.get('tipo_documento_ext', '')
                                id_ext = meta.get('id_documento_ext', '')
                                base_name = os.path.basename(src) if src else "Origem Desconhecida"

                                display_parts = []
                                if meta.get('origin') == 'temporary_attachment': display_parts.append(f"Anexo Temp '{orig_fname or base_name}'")
                                elif tipo == 'feedback': display_parts.append(f"Feedback/Análise {meta.get('feedback_timestamp', '')[:10]}")
                                elif tipo == 'feedback_attachment': display_parts.append(f"Anexo Feedback '{orig_fname or base_name}'")
                                elif tipo == 'texto_exemplo': display_parts.append(f"Exemplo Interno ({base_name})")
                                elif tipo_ext and tipo_ext not in ['DESCONHECIDO', 'ERRO_PARSE']: display_parts.append(f"Doc '{tipo_ext}'")
                                else: display_parts.append(f"Doc ({base_name})") # Nome base se tipo não for claro

                                if id_ext and id_ext != 'SEM_ID': display_parts.append(f"ID:{id_ext}")
                                display_parts.append(f"Chunk#{i+1}") # Adiciona número do chunk
                                fonte_display = " | ".join(display_parts) # Junta com separador claro

                                # Adiciona o chunk formatado ao contexto
                                contexto_formatado += f"\n--- INÍCIO: {fonte_display} ---\n{doc.page_content.strip()}\n--- FIM: {fonte_display} ---\n"
                            except Exception as e_fmt:
                                # Fallback se houver erro na formatação da fonte
                                st.warning(f"Erro ao formatar fonte para prompt: {e_fmt}")
                                contexto_formatado += f"\n--- INÍCIO: Fonte RAG #{i+1} ---\n{doc.page_content.strip()}\n--- FIM: Fonte RAG #{i+1} ---\n"
                    # Salva os metadados no estado da sessão para uso posterior (histórico, feedback)
                    st.session_state['fontes_usadas'] = fontes_usadas_meta_completas
                else:
                    st.session_state['fontes_usadas'] = [] # Garante que está vazio se não houver docs


                # Etapa 4: Gerar Resposta (Multi-chamada)
                model_name_selected = st.session_state.model_choice_selector
                model_context_limit = obter_limite_contexto(model_name_selected)
                model_output_limit = obter_limite_output(model_name_selected)
                safety_buffer = 250 # Aumenta buffer por segurança com prompts maiores
                respostas_geradas = [] # Lista para guardar as partes da resposta

                # Define quantas partes tentar gerar
                NUM_PARTES = 3

                if not docs_combinados: # --- Caso Sem RAG ---
                    status_text.info("⚠️ Nenhuma informação RAG encontrada. Gerando resposta geral...")
                    progress_bar.progress(0.75, text="Gerando resposta geral...")
                    prompt_geral = f"Você é CyberEvidence Oracle (IA jurídica/técnica). Sem informações específicas recuperadas, responda à pergunta abaixo baseado em conhecimento geral sobre o tópico, indicando claramente a ausência de fontes específicas consultadas para esta resposta.\n\nPergunta: {st.session_state.query}\n\nResposta Geral:"
                    try:
                        # Calcula max_tokens para resposta geral
                        prompt_tokens_geral = estimar_tokens(prompt_geral, model_name_selected)
                        available_completion_geral = model_context_limit - prompt_tokens_geral - safety_buffer
                        if available_completion_geral < 100: # Mínimo de tokens para gerar algo útil
                             st.error("Espaço insuficiente para gerar resposta geral com o modelo selecionado.")
                             raise ValueError("Contexto insuficiente para resposta geral.")

                        actual_max_tokens_geral = max(100, min(st.session_state.desired_max_tokens, available_completion_geral, model_output_limit))

                        resposta_gerada = criar_e_executar_llm(
                            prompt=prompt_geral,
                            model_name=model_name_selected,
                            temperature=st.session_state.temperature_slider,
                            max_tokens=actual_max_tokens_geral
                        )
                        if resposta_gerada and resposta_gerada.strip():
                            respostas_geradas.append(f"{resposta_gerada.strip()}\n\n---\n*Nota: Resposta baseada em conhecimento geral (sem consulta RAG específica). Modelo: {model_name_selected}*")
                        else:
                             respostas_geradas.append("[Erro: Resposta geral retornou vazia]")
                    except Exception as e_geral:
                        st.error(f"Erro ao gerar resposta geral: {e_geral}")
                        respostas_geradas.append(f"[Erro ao gerar resposta geral: {e_geral}]")

                else: # --- Caso Com RAG (Multi-chamada) ---
                    prompt_atual = ""
                    resposta_anterior_acumulada = "" # Acumula texto das respostas anteriores

                for parte_num in range(1, NUM_PARTES + 1):
                        progress_percentage = 0.55 + (parte_num * 0.30 / NUM_PARTES) # Distribui progresso
                        progress_bar.progress(progress_percentage, text=f"Gerando Resposta (Parte {parte_num}/{NUM_PARTES})...")

                        # Define o papel e as instruções específicas para cada parte
                        if parte_num == 1:
                            # == PARTE 1: FOCO EXCLUSIVO NA INTRODUÇÃO ==
                            papel_parte = "Introdução"
                            instrucoes_parte = f"""Instruções para a Parte 1 (Introdução):
1.  **Objetivo:** Elabore APENAS a introdução da resposta à pergunta original.
2.  **Contextualize:** Apresente brevemente o tema central da pergunta com base nos trechos RAG fornecidos.
3.  **Estrutura:** Esboce os principais pontos que serão abordados nas partes seguintes (desenvolvimento e conclusão), criando um roteiro claro para o leitor. NÃO desenvolva esses pontos aqui.
4.  **Fundamentação:** Baseie a contextualização e o esboço nos trechos RAG. Cite as fontes mais relevantes que dão o panorama geral (ex: "Os trechos de Anexo Temp 'relatorio.pdf' | Chunk#1 e Doc 'NormaXYZ' | Chunk#3 indicam que...").
5.  **Linguagem:** Clara, objetiva, didática e em Português (Brasil).
6.  **Tamanho:** Mantenha a introdução concisa e focada."""

                        elif parte_num == 2:
                            # == PARTE 2: FOCO EXCLUSIVO NO DESENVOLVIMENTO ==
                            papel_parte = "Desenvolvimento"
                            instrucoes_parte = f"""Instruções para a Parte 2 (Desenvolvimento):
1.  **Objetivo:** Elabore APENAS o desenvolvimento da resposta, detalhando os pontos esboçados na introdução (Parte 1).
2.  **Aprofundamento Didático:** Use extensivamente os trechos RAG fornecidos para:
    * **Explicar detalhadamente** cada ponto principal, definindo termos e clarificando conceitos como se estivesse ensinando.
    * **Fornecer exemplos concretos** ou analogias que ilustrem esses pontos, baseados nas fontes RAG ou em cenários plausíveis derivados delas, desde que deixe claro que a base não é o RAG.
    * **Descrever processos ou mecanismos** técnicos ou jurídicos mencionados de forma clara, sequencial e compreensível. Evite tópicos muito curtos. Desenvolva com trechos do RAG. Transcreva se for relevante. 
    * Analisar as implicações práticas ou teóricas dos pontos discutidos.
3.  **Conexão:** Garanta que o desenvolvimento flua logicamente a partir da introdução.
4.  **Fundamentação e Citação:** Cite PRECISAMENTENTE as fontes RAG para CADA informação ou argumento apresentado (ex: "Conforme Doc 'ManualY' | Chunk#15, o procedimento é..."). Use *apenas* os trechos fornecidos.
5.  **NÃO CONCLUA:** Evite sumarizar ou concluir a resposta nesta parte. Foque apenas na explicação e detalhamento.
6.  **Linguagem:** Mantenha o tom técnico quando necessário, mas sempre buscando a **máxima clareza, descrição e didática**."""

                        elif parte_num == 3:
                            # == PARTE 3: FOCO EXCLUSIVO NA CONCLUSÃO ==
                            papel_parte = "Conclusão"
                            instrucoes_parte = f"""Instruções para a Parte 3 (Conclusão):
1.  **Objetivo:** Elabore APENAS a conclusão da resposta, sintetizando os pontos discutidos nas partes anteriores (Introdução e Desenvolvimento).
2.  **Síntese:** Reafirme brevemente a ideia central da introdução e sumarize os principais argumentos/informações apresentados no desenvolvimento (Parte 2). NÃO introduza informações novas.
3.  **Fechamento:** Ofereça uma consideração final ou um fechamento coeso para a análise, respondendo diretamente à pergunta original com base na síntese feita.
4.  **Fundamentação:** A conclusão deve ser derivada logicamente das partes anteriores, que foram baseadas nos trechos RAG. Não é necessário repetir citações aqui, a menos que reforce um ponto chave da síntese.
5.  **Linguagem:** Clara, concisa e objetiva."""

                        else: # Fallback caso NUM_PARTES seja alterado
                             papel_parte = f"Parte {parte_num}"
                             instrucoes_parte = "Continue a resposta de forma coerente."

                        # Monta o prompt atual com as instruções específicas da parte
                        prompt_atual = f"""Você é um Professor de Direito Digital, um assistente de IA especialista em ensino e capacitação em Direito Digital e Legal StoryTelling. Sua tarefa é gerar uma resposta estruturada (Introdução, Desenvolvimento, Conclusão) à pergunta do usuário, baseando-se EXCLUSIVAMENTE nos trechos RAG fornecidos.

Siga estas diretrizes gerais:
- Responda em Português (Brasil).
- Use um tom claro, objetivo e **pedagógico.**
- Baseie TODAS as afirmações nos trechos RAG fornecidos. NÃO use conhecimento externo.
- Cite as fontes RAG relevantes conforme instruído para cada parte.
- **Explique conceitos técnicos ou jurídicos complexos de forma simples, como se estivesse ensinando.**
- **Use analogias ou exemplos práticos para ilustrar pontos abstratos sempre que possível.**
- **Defina termos-chave (jurídicos ou técnicos) importantes na primeira vez que aparecerem.**
- **Detalhe processos ou mecanismos de forma clara e sequencial.**
- Se encontrar informações conflitantes entre fontes gerais (documentos, anexos) e fontes de feedback/análise complementar, priorize a informação vinda do feedback.
                        
### PERGUNTA ORIGINAL DO USUÁRIO ###
{st.session_state.query}

### TRECHOS CONSULTADOS (Contexto RAG - Usar para TODAS as partes) ###
{contexto_formatado}
######################################

### RESPOSTA DAS PARTES ANTERIORES (se houver) ###
{resposta_anterior_acumulada if parte_num > 1 else "[Esta é a primeira parte, não há resposta anterior]"}
######################################

### SUA TAREFA AGORA (Parte {parte_num} de {NUM_PARTES}: {papel_parte}) ###
Siga RIGOROSAMENTE as instruções abaixo para gerar APENAS esta parte da resposta:
{instrucoes_parte}

### {papel_parte.upper()} (Parte {parte_num} de {NUM_PARTES}): ###""" # Header para o LLM começar a escrever

                        # --- O restante da lógica do loop (estimar tokens, chamar LLM, acumular resposta) permanece similar ---

                        # Estima tokens e calcula max_tokens para esta parte
                        prompt_tokens = estimar_tokens(prompt_atual, model_name_selected)
                        available_completion = model_context_limit - prompt_tokens - safety_buffer

                        if available_completion < 50: # Reduzido o mínimo para conclusão, mas ainda precisa de espaço
                            st.warning(f"Prompt RAG Parte {parte_num} ({papel_parte}) muito longo ({prompt_tokens} tk). Não é possível gerar esta parte.")
                            if parte_num == 1:
                                 st.error("Erro crítico: Prompt inicial (Introdução) excede o limite de contexto.")
                                 raise ValueError("Prompt inicial muito longo.")
                            break # Interrompe

                        # Ajusta max_tokens - talvez menos para intro/conclusão, mais para desenvolvimento?
                        # Pode-se refinar isso, mas usar o da UI é um bom começo.
                        desired_max_tokens_parte = st.session_state.desired_max_tokens
                        # Exemplo: reduzir tokens para conclusão
                        # if parte_num == NUM_PARTES:
                        #    desired_max_tokens_parte = max(100, min(desired_max_tokens_parte, 500))

                        actual_max_tokens = max(50, min(desired_max_tokens_parte, available_completion, model_output_limit))

                        status_text.info(f"Gerando {papel_parte} (~{prompt_tokens} tk, max {actual_max_tokens} tk)...")

                        try:
                            with st.spinner(f"⏳ Aguardando {model_name_selected} ({papel_parte})..."):
                                resposta_parte_atual = criar_e_executar_llm(
                                    prompt=prompt_atual,
                                    model_name=model_name_selected,
                                    temperature=st.session_state.temperature_slider,
                                    max_tokens=actual_max_tokens
                                )
                            resposta_limpa = resposta_parte_atual.strip() if resposta_parte_atual else ""

                            # Adiciona a resposta limpa à lista
                            if resposta_limpa:
                                respostas_geradas.append(resposta_limpa)
                                # Atualiza a resposta anterior ACUMULADA
                                # Para a próxima iteração, inclui um marcador claro da parte anterior
                                resposta_anterior_acumulada += f"\n\n--- {papel_parte.upper()} (Parte {parte_num} Gerada) ---\n" + resposta_limpa
                            else:
                                st.info(f"{papel_parte} (Parte {parte_num}) retornou vazia. Interrompendo.")
                                break # Interrompe se uma parte essencial retornar vazia

                        except Exception as e_llm_parte:
                            st.warning(f"Aviso: Erro ao gerar {papel_parte} (Parte {parte_num}): {e_llm_parte}")
                            break # Interrompe em caso de erro

                # Etapa 5: Combinar Respostas Finais e Salvar Histórico
                progress_bar.progress(0.95, text="Finalizando e formatando resposta...");

                if not respostas_geradas: # Se nenhuma parte foi gerada com sucesso
                    resposta_final_formatada = "[Erro: Nenhuma resposta válida pôde ser gerada]"
                    st.error("Falha na geração da resposta.")
                # Verifica se TODAS as partes esperadas (NUM_PARTES) foram geradas com sucesso
                elif len(respostas_geradas) == NUM_PARTES:
                    # Concatena Introdução, Desenvolvimento e Conclusão com parágrafos (junção natural)
                    # Usamos .strip() em cada parte para remover espaços extras antes/depois
                    resposta_final_formatada = f"{respostas_geradas[0].strip()}\n\n{respostas_geradas[1].strip()}\n\n{respostas_geradas[2].strip()}"
                    st.info(f"Resposta completa gerada com sucesso em {NUM_PARTES} partes (Intro/Dev/Concl).") # Mensagem de sucesso específica
                else: # Caso alguma parte tenha falhado ou retornado vazia (menos que NUM_PARTES)
                    resposta_final_formatada = "[Aviso: Falha ao gerar todas as partes estruturadas da resposta (Introdução/Desenvolvimento/Conclusão)]\n\n"
                    # Junta as partes que FORAM geradas, para análise do erro, com separador claro
                    resposta_final_formatada += "\n\n---\n".join([f"**Parte {i+1} Gerada:**\n{p.strip()}" for i, p in enumerate(respostas_geradas)])
                    st.warning(f"Apenas {len(respostas_geradas)} de {NUM_PARTES} partes estruturadas foram geradas com sucesso.")

                # Adiciona nota final sobre a geração (INDEPENDENTE de ter tido sucesso em todas as partes)
                num_partes_sucesso = len(respostas_geradas)
                total_chunks_usados = len(st.session_state.get('fontes_usadas', []))
                temp_usada = st.session_state.temperature_slider
                model_name_selected = st.session_state.model_choice_selector # Pega o nome do modelo usado

                resposta_final_formatada += f"\n\n---\n" # Separador antes da nota

                # Determina a nota sobre RAG baseado no contexto inicial
                if not docs_combinados: # Caso Sem RAG original
                     nota_final_rag = "*Nota: Resposta baseada em conhecimento geral (sem consulta RAG específica)."
                elif total_chunks_usados > 0:
                     nota_final_rag = f"*Consulta RAG: {total_chunks_usados} chunks utilizados."
                else: # Caso RAG foi tentado mas nenhum chunk foi usado/retornado efetivamente na combinação final
                     nota_final_rag = "*Nota: Consulta RAG realizada, mas nenhum chunk relevante encontrado ou utilizado para compor a resposta final."

                nota_final_modelo = f" Modelo: {model_name_selected}. Temp: {temp_usada}."
                # Ajusta a nota sobre as partes para ser mais clara
                if num_partes_sucesso == 1:
                    nota_final_partes = " Resposta gerada em 1 parte."
                elif num_partes_sucesso > 1:
                     nota_final_partes = f" Resposta gerada em {num_partes_sucesso} partes."
                else: # Caso 0 partes geradas com sucesso (já tratado no início, mas por segurança)
                     nota_final_partes = " Nenhuma parte da resposta gerada com sucesso."

                resposta_final_formatada += nota_final_rag + nota_final_modelo + nota_final_partes + "*"

                # Salva no histórico e no estado da sessão
                timestamp_resp = salvar_historico(st.session_state.query, resposta_final_formatada, st.session_state.get('fontes_usadas', []))
                st.session_state['resposta'] = resposta_final_formatada
                # Atualiza timestamp apenas se histórico foi salvo com sucesso
                if timestamp_resp: st.session_state['timestamp'] = timestamp_resp
                
                # ===============================================================
                # == FIM DO BLOCO DE ANÁLISE MULTI-RESPOSTA ==
                # ===============================================================

                # Etapa Final: Limpeza da UI e Arquivos Temporários
                progress_bar.progress(1.0); status_text.success("Análise concluída!")
                analysis_placeholder.empty() # Limpa mensagens de progresso/status

                # Limpa arquivos temporários salvos dos uploads
                if temp_paths_to_clean:
                    cleaned_count = 0
                    for temp_file in temp_paths_to_clean:
                        try:
                            if os.path.exists(temp_file) and os.path.isfile(temp_file):
                                os.remove(temp_file)
                                cleaned_count += 1
                        except Exception as e_clean:
                            st.warning(f"Erro ao limpar arquivo temporário {os.path.basename(temp_file)}: {e_clean}")
                    # st.info(f"Limpeza de {cleaned_count} arquivo(s) temporário(s) concluída.") # Opcional

                st.rerun() # Recarrega a UI para exibir a resposta final e o form de feedback

            except Exception as e_analysis: # Captura erro geral não tratado dentro do loop
                st.error(f"Erro inesperado durante o processo de análise: {e_analysis}")
                st.code(traceback.format_exc())
                try:
                    # Tenta limpar a UI de progresso mesmo em caso de erro
                    status_text.error(f"Erro durante a análise: {e_analysis}")
                    progress_bar.progress(1.0)
                    analysis_placeholder.empty()
                except Exception:
                    pass # Evita erro se placeholder já foi removido

                # Tenta limpar arquivos temporários mesmo em caso de erro na análise
                if temp_paths_to_clean:
                    st.info("Tentando limpar arquivos temporários após erro...")
                    for temp_file in temp_paths_to_clean:
                        try:
                            if os.path.exists(temp_file) and os.path.isfile(temp_file):
                                os.remove(temp_file)
                        except Exception as e_clean_on_error:
                             st.warning(f"Erro ao limpar temp {os.path.basename(temp_file)} após falha: {e_clean_on_error}")


    # --- Exibição Resposta e Feedback ---
    # Este bloco SÓ é exibido se houver uma resposta no estado da sessão
    if st.session_state.get('resposta'):
        st.markdown("---"); st.subheader("Resultado da Análise")
        # Exibe a resposta formatada (que pode conter markdown)
        st.markdown(st.session_state.resposta, unsafe_allow_html=False) # unsafe_allow_html=False é mais seguro

        # Expander para mostrar as fontes RAG consultadas (se houver)
        if st.session_state.get('fontes_usadas'):
            with st.expander("Fontes RAG Consultadas", expanded=False):
                fontes_display_set = set() # Usa um set para evitar listar a mesma fonte múltiplas vezes
                for meta in st.session_state.fontes_usadas:
                    display_name = "Fonte Desconhecida/Inválida"
                    if isinstance(meta, dict):
                         s_path = meta.get('source')
                         tipo = meta.get('tipo', 'Doc')
                         orig_fname = meta.get('original_filename') # Nome original (anexos/feedback)
                         tipo_ext = meta.get('tipo_documento_ext', '') # Extraído do nome (indexação)
                         id_ext = meta.get('id_documento_ext', '') # Extraído do nome (indexação)
                         origin = meta.get('origin', '') # Origem (temp, feedback)
                         ts_fb = meta.get('feedback_timestamp') # Timestamp do feedback

                         display_parts = []
                         # Tenta criar um nome descritivo
                         if origin == 'temporary_attachment': display_parts.append(f"Anexo Temp: {orig_fname or os.path.basename(s_path or '')}")
                         elif tipo == 'feedback': display_parts.append(f"Feedback/Análise ({ts_fb[:10] if ts_fb else 'Data N/A'})")
                         elif tipo == 'feedback_attachment': display_parts.append(f"Anexo Feedback: {orig_fname or os.path.basename(s_path or '')}")
                         elif tipo == 'texto_exemplo': display_parts.append("Exemplo Interno")
                         elif tipo_ext and tipo_ext not in ['DESCONHECIDO', 'ERRO_PARSE']: display_parts.append(f"Doc Tipo: {tipo_ext}")
                         else: display_parts.append(f"Doc: {os.path.basename(s_path or 'Nome N/A')}") # Nome base como fallback

                         if id_ext and id_ext != 'SEM_ID': display_parts.append(f"ID Ext: {id_ext}")

                         # Constrói o nome final para exibição
                         display_name = " | ".join(filter(None, display_parts)) # Junta partes não vazias
                         if not display_name: display_name = s_path if s_path else "Fonte sem nome" # Último fallback

                    fontes_display_set.add(display_name) # Adiciona ao set

                # Exibe as fontes únicas encontradas
                if fontes_display_set:
                    st.markdown("\n".join(f"- {f}" for f in sorted(list(fontes_display_set))))
                else:
                    st.markdown("- Nenhuma fonte RAG registrada ou metadados inválidos.")
        # Adiciona uma verificação para o caso de ter tido RAG mas a lista de fontes estar vazia
        elif 'docs_combinados' in locals() and docs_combinados and not st.session_state.get('fontes_usadas'):
             st.warning("Consulta usou RAG, mas as fontes não foram registradas corretamente no estado da sessão.")

        # --- Formulário de Feedback ---
        st.markdown("---"); st.subheader("Avaliação da Resposta")
        if st.session_state.get('feedback_submitted'):
             # Mensagem de sucesso após envio
             st.success("✅ Feedback enviado com sucesso! Obrigado pela sua contribuição.")
             # Botão para permitir novo feedback se necessário
             if st.button("Enviar Outro Feedback / Nova Avaliação"):
                 st.session_state['feedback_submitted'] = False
                 st.rerun() # Recarrega para mostrar o form novamente
        else:
            # Mostra o formulário se ainda não foi enviado
            with st.form("feedback_form", clear_on_submit=True): # Limpa campos após envio
                st.markdown("**Como você classificaria a utilidade e precisão da resposta gerada acima?**")
                # Opções de classificação
                f_class = st.radio(
                    "Classificação:",
                    options=["⭐⭐⭐⭐⭐ Ótima", "⭐⭐⭐⭐ Boa", "⭐⭐⭐ Razoável", "⭐⭐ Ruim", "⭐ Inútil"],
                    index=None, # Nenhum selecionado por padrão
                    horizontal=True,
                    key="f_radio"
                )
                # Campo para comentários
                f_comm = st.text_area(
                    "Comentário / Análise Complementar (Opcional):",
                    key="f_comment",
                    placeholder="Detalhe sua avaliação. Aponte erros ou imprecisões. Sugira melhorias ou informações faltantes. Sua análise pode ser indexada para consultas futuras.",
                    height=150
                )
                # Uploader para anexos de feedback
                f_files = st.file_uploader(
                    "Anexar arquivos para complementar/justificar (Opcional, PDF/DOCX/TXT):",
                    accept_multiple_files=True,
                    key="f_uploader",
                    type=['pdf', 'docx', 'txt'], 
                    help="Anexe documentos que complementem sua análise, corrijam a resposta ou forneçam contexto adicional. O conteúdo será indexado."
                )
                # Botão de envio do formulário
                submitted_feedback = st.form_submit_button("✔️ Enviar Avaliação")

                if submitted_feedback:
                    # Validação: Classificação é obrigatória
                    if not f_class:
                        st.warning("⚠️ Por favor, selecione uma classificação para a resposta.")
                    else:
                        # Processa e salva o feedback
                        with st.spinner("Processando e salvando seu feedback..."):
                             # Pega a query e resposta do estado da sessão associadas a este feedback
                             q_orig = st.session_state.current_query_for_feedback
                             r_aval = st.session_state.resposta
                             # Usa o timestamp da resposta original, se disponível, senão o atual
                             ts_fb = st.session_state.timestamp or datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                             # Chama a função para salvar (CSV, RAG do comentário, RAG dos anexos)
                             salvar_feedback(q_orig, r_aval, f_comm, f_class, f_files or [], ts_fb)

                             st.session_state['feedback_submitted'] = True # Marca como enviado
                             # Limpa cache da auditoria para forçar recarregamento se o usuário visualizar
                             if 'df_feedback_audit' in st.session_state: del st.session_state['df_feedback_audit']
                             st.session_state['feedback_loaded'] = False
                        # Recarrega a página para mostrar a mensagem de sucesso e limpar o form
                        st.rerun()

# Ponto de entrada principal da aplicação Streamlit
if __name__ == "__main__":
    # Chama a função principal que contém a lógica da UI e da aplicação
    main()